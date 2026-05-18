"""
Genera il documento Word "Guida_Firebase_Silvestre.docx" con stato attuale
del backend Firebase e cosa serve fare per renderlo pienamente operativo.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date

ORANGE = RGBColor(0xF4, 0x75, 0x21)
DARK = RGBColor(0x2B, 0x2B, 0x2B)
GREY = RGBColor(0x7A, 0x7A, 0x7A)
RED = RGBColor(0xD6, 0x45, 0x45)
GREEN = RGBColor(0x3D, 0xA3, 0x5D)


def add_heading(doc, text, level=1, color=ORANGE):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = "Calibri"


def add_p(doc, text, bold=False, italic=False, color=DARK, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    return p


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(it)
        run.font.size = Pt(11)
        run.font.name = "Calibri"


def add_numbered(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(it)
        run.font.size = Pt(11)
        run.font.name = "Calibri"


def add_checklist(doc, items):
    for it in items:
        p = doc.add_paragraph()
        chk = p.add_run("\u2610  ")
        chk.font.size = Pt(13)
        chk.bold = True
        chk.font.color.rgb = ORANGE
        txt = p.add_run(it)
        txt.font.size = Pt(11)
        txt.font.name = "Calibri"


def add_done(doc, text):
    p = doc.add_paragraph()
    chk = p.add_run("\u2705  ")
    chk.font.size = Pt(13)
    txt = p.add_run(text)
    txt.font.size = Pt(11)
    txt.font.name = "Calibri"
    txt.font.color.rgb = GREEN
    txt.bold = True


def add_todo(doc, text):
    p = doc.add_paragraph()
    chk = p.add_run("\u23f3  ")
    chk.font.size = Pt(13)
    txt = p.add_run(text)
    txt.font.size = Pt(11)
    txt.font.name = "Calibri"
    txt.bold = True


def add_warning(doc, text):
    p = doc.add_paragraph()
    run = p.add_run("\u26a0  " + text)
    run.bold = True
    run.font.color.rgb = RED
    run.font.size = Pt(11)
    run.font.name = "Calibri"


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F47521")
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
    for r, row in enumerate(rows, start=1):
        cells = t.rows[r].cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10.5)
            run.font.name = "Calibri"


# ============================================================================
doc = Document()
styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

# --- COPERTINA ---
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Guida Firebase")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = ORANGE
r.font.name = "Calibri"

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Backend dell'app Silvestre Fotoservizi")
r.font.size = Pt(15)
r.font.color.rgb = DARK
r.font.name = "Calibri"

dt = doc.add_paragraph()
dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = dt.add_run(f"Versione del {date.today().strftime('%d/%m/%Y')}")
r.font.size = Pt(10)
r.font.color.rgb = GREY
r.font.name = "Calibri"

# --- INTRO ---
add_heading(doc, "Cos'\u00e8 e perch\u00e9 lo usiamo", level=1)
add_p(
    doc,
    "Firebase \u00e8 la piattaforma di Google che fa da \"cervello\" centrale dell'app: "
    "tiene i dati (ordini, profili, foto), gestisce i login degli utenti, sincronizza "
    "le informazioni tra app cliente e app operatore in tempo reale, e manda le notifiche. "
    "Senza Firebase ogni installazione dell'app sarebbe isolata e i dati svanirebbero al refresh.",
)

add_p(doc, "I servizi che useremo:")
add_bullets(
    doc,
    [
        "Authentication \u2014 login con email/password (in futuro anche Google/Apple).",
        "Cloud Firestore \u2014 database documentale per ordini, profili, prodotti.",
        "Cloud Storage \u2014 dove vivono le foto caricate dai clienti.",
        "Cloud Functions \u2014 piccoli script che girano sui server di Google per eventi (es. invio notifica all'operatore quando arriva un nuovo ordine).",
        "Cloud Messaging \u2014 push notification al cliente e all'operatore.",
    ],
)

doc.add_page_break()

# --- IDENTITA PROGETTO ---
add_heading(doc, "Identit\u00e0 del progetto Firebase", level=1)
add_table(
    doc,
    ["Campo", "Valore"],
    [
        ["Project ID", "silvestre-fotoservizi"],
        ["Project Number", "1029797648749"],
        ["Display Name", "Silvestre Fotoservizi"],
        ["Console URL", "https://console.firebase.google.com/project/silvestre-fotoservizi/overview"],
        ["Web App ID", "1:1029797648749:web:31017e2ef142a3a46b2264"],
        ["Auth Domain", "silvestre-fotoservizi.firebaseapp.com"],
        ["Storage Bucket", "silvestre-fotoservizi.firebasestorage.app"],
    ],
)
add_p(
    doc,
    "Le chiavi (apiKey, appId, messagingSenderId) sono salvate in lib/firebase_options.dart "
    "dell'app Flutter. Sono CLIENT-SIDE e per design pubbliche; la sicurezza vera \u00e8 "
    "nelle Security Rules.",
    italic=True,
)

doc.add_page_break()

# --- STATO ATTUALE ---
add_heading(doc, "Stato attuale (cosa \u00e8 gi\u00e0 fatto)", level=1)

add_done(doc, "Progetto GCP creato (Silvestre Fotoservizi).")
add_done(doc, "Firebase aggiunto al progetto e ToS accettati.")
add_done(doc, "Firestore Database provisionato (regione default).")
add_done(doc, "Firestore Security Rules scritte e deployate (firebase/firestore.rules).")
add_done(doc, "Indici Firestore configurati (firebase/firestore.indexes.json).")
add_done(doc, "Storage Security Rules scritte (firebase/storage.rules) \u2014 pronte ma in attesa che si abiliti Storage.")
add_done(doc, "Web App creata in Firebase Console.")
add_done(doc, "lib/firebase_options.dart generato nell'app Flutter.")
add_done(doc, "Shortcut a Firebase Console creato in cartella principale.")

add_p(doc, "")

add_heading(doc, "Cosa manca (da completare insieme)", level=2)
add_todo(doc, "Abilitare Cloud Storage via Console (1 click di setup).")
add_todo(doc, "Abilitare Authentication > Email/Password come provider iniziale.")
add_todo(doc, "Aggiungere FirebaseFlutter dependencies a pubspec.yaml e wire-up.")
add_todo(doc, "Migrare AuthState/OrdersState/CartState da mock a Firestore.")
add_todo(doc, "Creare app Android e iOS in Firebase Console (per build mobile).")
add_todo(doc, "Configurare Cloud Functions per notifiche e moderazione contenuti.")
add_todo(doc, "Setup billing Spark (free) o passaggio a Blaze (pay-as-you-go).")

doc.add_page_break()

# --- ABILITARE STORAGE ---
add_heading(doc, "Abilitare Cloud Storage (1 click)", level=1)
add_numbered(
    doc,
    [
        "Apri lo shortcut \"Firebase Console.url\" in cartella SilvestreApp.",
        "Nel menu a sinistra: Build > Storage.",
        "Clicca \"Get Started\" / \"Inizia\".",
        "Modalit\u00e0 di sicurezza: scegli \"Start in production mode\" (le nostre regole vere sono nel file storage.rules).",
        "Regione: scegli europe-west3 (Francoforte) per essere conformi GDPR.",
        "Conferma. Aspetta 30 secondi.",
        "Quando vedi la dashboard \"Files\", torna in chat e dimmi che hai finito.",
    ],
)
add_p(
    doc,
    "Dopo questo passo io potr\u00f2 deployare le regole Storage da CLI con un singolo comando "
    "(gi\u00e0 pronto in script).",
    italic=True,
)

doc.add_page_break()

# --- ABILITARE AUTH ---
add_heading(doc, "Abilitare Authentication (email/password)", level=1)
add_numbered(
    doc,
    [
        "Firebase Console > Authentication > Get Started.",
        "Scheda \"Sign-in method\".",
        "Clicca \"Email/Password\" \u2192 toggle on \u2192 Save.",
        "Successivamente potrai abilitare anche \"Google\" e \"Apple\" come provider opzionali.",
    ],
)

doc.add_page_break()

# --- SECURITY RULES ---
add_heading(doc, "Security Rules \u2014 come funzionano", level=1)
add_p(
    doc,
    "Le regole vivono in firebase/firestore.rules e firebase/storage.rules. Sono il "
    "VERO meccanismo di sicurezza: anche se la chiave API \u00e8 pubblica, le regole "
    "impediscono accessi non autorizzati.",
)

add_heading(doc, "Modello dei ruoli", level=2)
add_bullets(
    doc,
    [
        "customer \u2014 utente normale (default per chi si registra dall'app).",
        "staff \u2014 dipendente del negozio (vede tutti gli ordini, pu\u00f2 cambiarne lo stato).",
        "admin \u2014 proprietario (CRUD prodotti, gestione staff, cancellazione ordini).",
    ],
)

add_heading(doc, "Permessi sintetici", level=2)
add_table(
    doc,
    ["Collezione", "Customer", "Staff", "Admin"],
    [
        ["users/{me}", "R/W del proprio", "R di tutti", "R/W di tutti"],
        ["products", "R", "R", "R/W"],
        ["templates", "R", "R", "R/W"],
        ["settings", "R", "R", "R/W"],
        ["orders (own)", "Create + R", "R + Update status", "R/W/Delete"],
        ["designs (own)", "R/W del proprio", "R", "R"],
        ["uploads/{me} (Storage)", "R/W del proprio (max 20MB, solo immagini)", "R", "R"],
    ],
)

add_heading(doc, "Modificare le regole", level=2)
add_numbered(
    doc,
    [
        "Modifica firebase/firestore.rules o firebase/storage.rules.",
        "Da terminale, nella cartella firebase, esegui:",
        "  firebase deploy --only firestore:rules --token \"$FIREBASE_TOKEN\"",
        "  firebase deploy --only storage --token \"$FIREBASE_TOKEN\"",
        "Le regole sono attive entro 10 secondi globalmente.",
    ],
)

doc.add_page_break()

# --- CREARE APP MOBILE ---
add_heading(doc, "Creare App Android e iOS in Firebase", level=1)
add_p(
    doc,
    "La Web App \u00e8 gi\u00e0 creata. Per gli store mobile servono app dedicate (questo aggiunge "
    "i file google-services.json per Android e GoogleService-Info.plist per iOS).",
)

add_heading(doc, "Android", level=2)
add_numbered(
    doc,
    [
        "Firebase Console > Project Settings (rotella ingranaggio) > tab General.",
        "Scroll fino a \"Your apps\" \u2192 icona Android.",
        "Package name: com.silvestrefotoservizi.app",
        "App nickname: \"Silvestre Android\".",
        "SHA-1 (opzionale per ora, obbligatorio per Google Sign-In e App Links).",
        "Register app \u2192 scarica google-services.json.",
        "Copia il file in: SilvestreApp/app/android/app/google-services.json",
    ],
)

add_heading(doc, "iOS", level=2)
add_numbered(
    doc,
    [
        "Firebase Console > Project Settings > tab General > icona iOS.",
        "Bundle ID: com.silvestrefotoservizi.app",
        "App nickname: \"Silvestre iOS\".",
        "App Store ID (lascia vuoto finch\u00e9 non pubblichi).",
        "Register app \u2192 scarica GoogleService-Info.plist.",
        "Copia il file in: SilvestreApp/app/ios/Runner/GoogleService-Info.plist",
    ],
)

add_heading(doc, "Aggiornare firebase_options.dart", level=2)
add_p(
    doc,
    "Dopo aver creato le app Android e iOS in Console, lancia da terminale (nella cartella app):",
)
add_p(doc, "  dart pub global activate flutterfire_cli", italic=True)
add_p(doc, "  flutterfire configure --project=silvestre-fotoservizi", italic=True)
add_p(doc,
      "Il comando rigenera lib/firebase_options.dart con le chiavi corrette per ogni piattaforma.")

doc.add_page_break()

# --- BILLING ---
add_heading(doc, "Billing \u2014 Spark vs Blaze", level=1)
add_p(
    doc,
    "Il progetto parte sul piano Spark (gratuito). Limiti fondamentali del piano Spark:",
)
add_table(
    doc,
    ["Servizio", "Limite Spark (gratis)", "Quando passare a Blaze"],
    [
        ["Authentication", "50.000 MAU", "Mai superato in negozio piccolo"],
        ["Firestore reads", "50.000/giorno", "Sopra ~3.000 ordini visualizzati/giorno"],
        ["Firestore writes", "20.000/giorno", "Sopra ~500-1.000 ordini scritti/giorno"],
        ["Firestore storage", "1 GiB", "Sopra ~500.000 righe ordini"],
        ["Storage download", "5 GiB/giorno", "Sopra ~3.000 foto scaricate/giorno"],
        ["Storage capacity", "5 GiB totale", "Sopra ~2.500 foto da 2MB"],
        ["Cloud Functions", "Non disponibile su Spark", "Subito che servono notifiche server-side"],
        ["Cloud Messaging", "Illimitato", "Mai"],
    ],
)
add_warning(
    doc,
    "Cloud Functions richiede Blaze (pay-as-you-go). Per le notifiche operatore servono Functions. "
    "Il primo passaggio costa pochi euro/mese in pratica per un negozio piccolo (~2-10\u20ac).",
)

add_heading(doc, "Passare a Blaze", level=2)
add_numbered(
    doc,
    [
        "Firebase Console > Settings > Usage and billing > Details and settings > Modify plan.",
        "Scegli Blaze (pay-as-you-go).",
        "Collega carta di credito (Google addebita solo l'eccedenza oltre le quote gratuite).",
        "IMPORTANTE: imposta un Budget Alert a 10\u20ac/mese in Google Cloud Console per essere avvisato.",
    ],
)

doc.add_page_break()

# --- INTEGRAZIONE FLUTTER ---
add_heading(doc, "Integrazione nell'app Flutter (prossimo step)", level=1)
add_p(doc, "Quando sar\u00e0 il momento di collegare l'app a Firebase, i passi tecnici sono:")
add_numbered(
    doc,
    [
        "Aggiungere a pubspec.yaml: firebase_core, firebase_auth, cloud_firestore, firebase_storage, firebase_messaging.",
        "flutter pub get.",
        "In main.dart: WidgetsFlutterBinding.ensureInitialized() + await Firebase.initializeApp(options: DefaultFirebaseOptions.currentPlatform).",
        "Sostituire AuthState mock con un AuthState che usa FirebaseAuth.instance.",
        "Sostituire OrdersState mock con un OrdersState che ascolta FirebaseFirestore.instance.collection('orders')...",
        "Sostituire CartState mock con un cart che si salva su 'users/{uid}/cart'.",
        "Aggiungere caricamento foto via FirebaseStorage.instance.",
    ],
)
add_p(
    doc,
    "Tutto questo \u00e8 lavoro mio \u2014 ti basta dare il go.",
    italic=True,
)

doc.add_page_break()

# --- TROUBLESHOOTING ---
add_heading(doc, "Troubleshooting comune", level=1)

add_heading(doc, "\"Permission denied\" su una collezione", level=2)
add_bullets(
    doc,
    [
        "L'utente non \u00e8 loggato \u2192 verifica request.auth nelle Rules.",
        "Il documento ha userId diverso dall'utente \u2192 controllo isOwner.",
        "Le Rules non sono ancora deployate \u2192 firebase deploy --only firestore:rules.",
    ],
)

add_heading(doc, "\"Quota exceeded\"", level=2)
add_bullets(
    doc,
    [
        "Stai sul piano Spark e hai superato i limiti \u2192 passa a Blaze.",
        "C'\u00e8 un loop di query inefficiente \u2192 controlla Crashlytics + Console > Usage.",
    ],
)

add_heading(doc, "\"App not authorized to use Firebase\"", level=2)
add_bullets(
    doc,
    [
        "Hai cambiato Package name o Bundle ID dopo la registrazione \u2192 rigenera l'app in Console.",
        "Hai dimenticato di copiare google-services.json o GoogleService-Info.plist nel progetto.",
    ],
)

# --- LINK UTILI ---
add_heading(doc, "Link utili", level=1)
add_bullets(
    doc,
    [
        "Console: https://console.firebase.google.com/project/silvestre-fotoservizi",
        "Documentazione: https://firebase.google.com/docs",
        "Status Google Cloud: https://status.cloud.google.com",
        "Prezzi: https://firebase.google.com/pricing",
        "FlutterFire docs: https://firebase.flutter.dev",
    ],
)

# --- SALVA ---
import os
out = os.path.join(os.path.dirname(__file__), "Guida_Firebase_Silvestre.docx")
doc.save(out)
print(f"OK: {out}")
