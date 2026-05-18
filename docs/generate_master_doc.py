"""
generate_master_doc.py — unisce TUTTI i .docx in un unico MANUALE_GENERALE_Silvestre.docx.

Workflow:
1. Rigenera ogni singolo .docx invocando i generate_*.py
2. Apre il manuale generale come base
3. Aggiunge in coda: Conformità Legale, Guida Firebase, Pubblicazione App,
   Deploy FCM Mobile — ciascuno preceduto da una pagina di copertina.

Risultato: un unico documento master con tutte le sezioni, navigabile via TOC.
"""
import os
import subprocess
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxcompose.composer import Composer

DOCS_DIR = os.path.dirname(os.path.abspath(__file__))
MASTER_PATH = os.path.join(DOCS_DIR, "MANUALE_GENERALE_Silvestre.docx")

# Ordine in cui includere i sotto-documenti dopo il manuale generale
APPENDICES = [
    ("Conformita_Legale_Silvestre.docx",
     "APPENDICE A — Conformità Legale e GDPR"),
    ("Guida_Firebase_Silvestre.docx",
     "APPENDICE B — Guida Firebase (backend dati)"),
    ("Pubblicazione_App_Silvestre.docx",
     "APPENDICE C — Pubblicazione App Store / Play Store"),
    ("Deploy_FCM_Mobile_Functions.docx",
     "APPENDICE D — Deploy FCM e Cloud Functions"),
]

GENERATORS = [
    "generate_manuale_generale.py",
    "generate_legal_doc.py",
    "generate_firebase_doc.py",
    "generate_launch_doc.py",
    "generate_fcm_mobile_deploy_doc.py",
]


def regen_all():
    """Rigenera ogni .docx individuale."""
    for g in GENERATORS:
        path = os.path.join(DOCS_DIR, g)
        if not os.path.exists(path):
            print(f"  SKIP: {g} non trovato")
            continue
        print(f"  Rigenero {g}...")
        r = subprocess.run([sys.executable, path], capture_output=True, text=True)
        if r.returncode != 0:
            print(f"    ERRORE: {r.stderr[:300]}")
        else:
            # Print last line of output
            last = (r.stdout.strip().splitlines() or [""])[-1]
            print(f"    {last}")


def add_section_cover(doc, title):
    """Aggiunge una pagina di copertina per l'appendice."""
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n\n\n")
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run(title)
    run.bold = True
    run.font.size = doc.styles["Heading 1"].font.size or None
    doc.add_page_break()


def main():
    print("=== STEP 1: rigenero tutti i .docx individuali ===")
    regen_all()

    print(f"\n=== STEP 2: compongo {MASTER_PATH} ===")
    base_path = os.path.join(DOCS_DIR, "MANUALE_GENERALE_Silvestre.docx")
    if not os.path.exists(base_path):
        print(f"ERRORE: {base_path} non esiste")
        sys.exit(1)

    base = Document(base_path)
    composer = Composer(base)

    for fname, title in APPENDICES:
        path = os.path.join(DOCS_DIR, fname)
        if not os.path.exists(path):
            print(f"  SKIP: {fname} non trovato")
            continue
        print(f"  Aggiungo: {title}")

        # Cover page con titolo appendice
        cover_doc = Document()
        cover_doc.add_page_break()
        p = cover_doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for _ in range(8):
            p.add_run("\n")
        p2 = cover_doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p2.add_run(title)
        run.bold = True
        run.font.size = None  # use default
        from docx.shared import Pt
        run.font.size = Pt(28)
        composer.append(cover_doc)

        # Contenuto appendice
        appendix = Document(path)
        composer.append(appendix)

    composer.save(MASTER_PATH)
    print(f"\nOK: {MASTER_PATH}")
    size_kb = os.path.getsize(MASTER_PATH) / 1024
    print(f"Dimensione: {size_kb:.0f} KB")


if __name__ == "__main__":
    main()
