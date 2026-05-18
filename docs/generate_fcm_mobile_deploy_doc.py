"""
Genera "Deploy_FCM_Mobile_Functions.docx" con i passi finali pre-rilascio:
- Setup FCM (VAPID web + APNs iOS)
- Registrazione app mobile Android/iOS in Firebase
- Deploy Cloud Functions (richiede Blaze)
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date

ORANGE = RGBColor(0xF4, 0x75, 0x21)
DARK = RGBColor(0x2B, 0x2B, 0x2B)
GREY = RGBColor(0x7A, 0x7A, 0x7A)
RED = RGBColor(0xD6, 0x45, 0x45)
GREEN = RGBColor(0x3D, 0xA3, 0x5D)


def h(doc, t, level=1, color=ORANGE):
    p = doc.add_heading(t, level=level)
    for r in p.runs:
        r.font.color.rgb = color
        r.font.name = "Calibri"


def p(doc, t, bold=False):
    par = doc.add_paragraph()
    run = par.add_run(t)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def num(doc, items):
    for it in items:
        par = doc.add_paragraph(style="List Number")
        run = par.add_run(it)
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def bul(doc, items):
    for it in items:
        par = doc.add_paragraph(style="List Bullet")
        run = par.add_run(it)
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def warn(doc, t):
    par = doc.add_paragraph()
    r = par.add_run("[!] " + t)
    r.bold = True
    r.font.color.rgb = RED
    r.font.size = Pt(11)
    r.font.name = "Calibri"


def code(doc, t):
    par = doc.add_paragraph()
    r = par.add_run(t)
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x36, 0x36, 0x36)


doc = Document()
styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Deploy: FCM + Mobile + Functions")
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = ORANGE
r.font.name = "Calibri"

dt = doc.add_paragraph()
dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = dt.add_run(f"Versione del {date.today().strftime('%d/%m/%Y')}")
r.font.size = Pt(10)
r.font.color.rgb = GREY
r.font.name = "Calibri"

p(doc,
  "Passi finali per portare l'app dallo stato 'web demo' a 'pubblicabile su store mobile' "
  "con push notification e backend automatizzato. Tutto il codice e' gia' scritto: questi "
  "sono i pochi click in console che mancano.", bold=False)

# ----------- FCM Web Push -----------
h(doc, "1. Web Push (VAPID key)", level=1)
p(doc, "L'app Flutter web e' gia' configurata col pacchetto firebase_messaging. Manca solo "
       "la chiave VAPID per consentire le push browser.")
num(doc, [
    "Firebase Console: console.firebase.google.com/project/silvestre-fotoservizi",
    "Project Settings (rotella in alto) -> tab Cloud Messaging",
    "Sezione 'Web Push certificates' -> 'Generate key pair'",
    "Copia la stringa 'Key pair' generata (e.g. BABCDEF... lunga ~88 char)",
    "Apri app/lib/services/push_notifications_service.dart",
    "Sostituisci 'static const String? webVapidKey = null;' con la tua chiave:",
])
code(doc, "static const String? webVapidKey = 'BABCDEF...la_tua_chiave';")
num(doc, [
    "Salva, rilancia il dev server.",
    "Al login del cliente il browser chiedera' il permesso notifiche.",
])

# ----------- iOS APNs -----------
h(doc, "2. iOS push (APNs)", level=1)
p(doc, "Solo per build iOS. Richiede Apple Developer account.")
num(doc, [
    "Apple Developer Center -> Certificates, Identifiers & Profiles",
    "Crea APNs Authentication Key (file .p8) — segui guida ufficiale",
    "Firebase Console -> Project Settings -> Cloud Messaging -> Apple app configuration",
    "Carica il .p8, inserisci Key ID + Team ID",
    "Fine: le push iOS arrivano automaticamente.",
])

# ----------- Android FCM -----------
h(doc, "3. Android push", level=1)
p(doc, "Funziona automaticamente una volta che hai registrato l'app Android (sez. 5).")

# ----------- Cloud Functions deploy -----------
h(doc, "4. Cloud Functions deploy", level=1)
warn(doc, "Richiede piano Blaze (pay-as-you-go con carta di credito).")
p(doc, "Le 3 funzioni sono gia' scritte in firebase/functions/index.js:")
bul(doc, [
    "onOrderCreated -> notifica tutti gli operatori al nuovo ordine",
    "onOrderStatusChange -> notifica cliente al cambio stato",
    "scheduledPhotoCleanup -> ogni notte 03:00 cancella foto >30gg post-ritiro",
])

p(doc, "Setup Blaze:")
num(doc, [
    "Firebase Console -> rotella ingranaggio -> Usage and billing -> Details and settings -> Modify plan",
    "Scegli Blaze",
    "Collega carta",
    "Imposta Budget Alert a 5 EUR/mese in Google Cloud Console -> Billing -> Budgets",
])

p(doc, "Setup Cloudinary API key per cleanup:")
num(doc, [
    "Cloudinary dashboard -> Settings -> Access Keys -> copia API Key + Secret",
    "In firebase/ esegui:",
])
code(doc, 'firebase functions:secrets:set CLOUDINARY_API_KEY')
code(doc, 'firebase functions:secrets:set CLOUDINARY_API_SECRET')
p(doc, "(incolli i valori quando richiesto, restano cifrati su Google Secret Manager)")

p(doc, "Install dependencies + deploy:")
code(doc, 'cd firebase/functions')
code(doc, 'npm install')
code(doc, 'cd ..')
code(doc, 'firebase deploy --only functions --token "$FIREBASE_TOKEN"')

p(doc, "Verifica:")
num(doc, [
    "Firebase Console -> Functions -> dovresti vedere 3 funzioni 'deployed'",
    "Crea un ordine test dall'app cliente",
    "Apri Logs della funzione onOrderCreated -> vedi 'New order ...' loggato",
])

# ----------- Mobile App Registration -----------
h(doc, "5. Registrazione app Android in Firebase", level=1)
num(doc, [
    "Firebase Console -> Project Settings -> tab General -> scroll a 'Your apps'",
    "Icona Android -> Add app",
    "Package name: com.silvestrefotoservizi.app",
    "App nickname: 'Silvestre Android'",
    "SHA-1: opzionale per ora (lascia vuoto). Necessario per Google Sign-In.",
    "Register app -> scarica google-services.json",
    "Sposta il file in: SilvestreApp/app/android/app/google-services.json",
    "Apri app/android/app/build.gradle e aggiungi in fondo:",
])
code(doc, "apply plugin: 'com.google.gms.google-services'")
num(doc, [
    "Apri app/android/build.gradle, in plugins{} aggiungi:",
])
code(doc, 'id "com.google.gms.google-services" version "4.4.2" apply false')

h(doc, "6. Registrazione app iOS in Firebase", level=1)
num(doc, [
    "Firebase Console -> Project Settings -> tab General -> Add app -> iOS",
    "Bundle ID: com.silvestrefotoservizi.app",
    "App nickname: 'Silvestre iOS'",
    "Register app -> scarica GoogleService-Info.plist",
    "Sposta in: SilvestreApp/app/ios/Runner/GoogleService-Info.plist",
    "Apri il file ios/Runner.xcworkspace in Xcode (sul tuo Mac)",
    "Drag-and-drop GoogleService-Info.plist sotto il gruppo Runner",
    "Selezionalo, in inspector destra spunta 'Runner' target",
])

# ----------- Rigenera firebase_options.dart -----------
h(doc, "7. Rigenera firebase_options.dart con tutte le piattaforme", level=1)
p(doc, "Dopo aver registrato Android e iOS in Console, lancia FlutterFire CLI per regenerare automaticamente:")
code(doc, "dart pub global activate flutterfire_cli")
code(doc, "cd app")
code(doc, "flutterfire configure --project=silvestre-fotoservizi")
p(doc, "Il comando rigenera lib/firebase_options.dart con le 3 piattaforme corrette (web + android + iOS).")

# ----------- Final checklist -----------
h(doc, "8. Checklist finale", level=1)
ck = [
    "VAPID key web sostituita in push_notifications_service.dart",
    "Apple Developer attivo (99 EUR/anno) e APNs .p8 caricato in Firebase",
    "App Android registrata in Firebase + google-services.json in app/android/app/",
    "App iOS registrata in Firebase + GoogleService-Info.plist in app/ios/Runner/",
    "flutterfire configure eseguito (firebase_options.dart aggiornato per 3 piattaforme)",
    "Blaze plan attivato + Budget Alert 5 EUR/mese",
    "Cloudinary API Key + Secret in Firebase Secrets",
    "firebase deploy --only functions completato senza errori",
    "Test: crea ordine cliente -> arriva push operatore",
    "Test: operatore cambia stato -> arriva push cliente",
    "Apple App Store Connect creato (per build iOS)",
    "Google Play Console creato (per build Android)",
]
for c in ck:
    par = doc.add_paragraph()
    chk = par.add_run("[ ]  ")
    chk.bold = True
    chk.font.size = Pt(13)
    chk.font.color.rgb = ORANGE
    t = par.add_run(c)
    t.font.size = Pt(11)
    t.font.name = "Calibri"

import os
out = os.path.join(os.path.dirname(__file__), "Deploy_FCM_Mobile_Functions.docx")
doc.save(out)
print(f"OK: {out}")
