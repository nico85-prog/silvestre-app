"""
MANUALE_GENERALE_Silvestre.docx — guida compatta e scorrevole.
Entry-point: contiene tutto l'essenziale + rimanda agli altri docs per il dettaglio.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date

ORANGE = RGBColor(0xF4, 0x75, 0x21)
DARK = RGBColor(0x2B, 0x2B, 0x2B)
GREY = RGBColor(0x7A, 0x7A, 0x7A)
RED = RGBColor(0xD6, 0x45, 0x45)
GREEN = RGBColor(0x3D, 0xA3, 0x5D)
BLUE = RGBColor(0x1F, 0x77, 0xB4)


def h(doc, t, level=1, color=ORANGE):
    par = doc.add_heading(t, level=level)
    for r in par.runs:
        r.font.color.rgb = color
        r.font.name = "Calibri"


def p(doc, t, bold=False, italic=False, color=DARK, size=11):
    par = doc.add_paragraph()
    r = par.add_run(t)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = color
    r.font.size = Pt(size)
    r.font.name = "Calibri"


def bul(doc, items):
    for it in items:
        par = doc.add_paragraph(style="List Bullet")
        r = par.add_run(it)
        r.font.size = Pt(11)
        r.font.name = "Calibri"


def num(doc, items):
    for it in items:
        par = doc.add_paragraph(style="List Number")
        r = par.add_run(it)
        r.font.size = Pt(11)
        r.font.name = "Calibri"


def warn(doc, t):
    par = doc.add_paragraph()
    r = par.add_run("\u26a0  " + t)
    r.bold = True
    r.font.color.rgb = RED
    r.font.size = Pt(11)
    r.font.name = "Calibri"


def tip(doc, t):
    par = doc.add_paragraph()
    r = par.add_run("\U0001f4a1  " + t)
    r.font.color.rgb = GREEN
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = "Calibri"


def table(doc, headers, rows, col_widths_inches=None):
    """col_widths_inches: lista di larghezze (in pollici) per ciascuna colonna.
    Se None, usa larghezze automatiche."""
    from docx.shared import Inches
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.autofit = False
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = ""
        par = hdr[i].paragraphs[0]
        run = par.add_run(htxt)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F47521")
        hdr[i]._tc.get_or_add_tcPr().append(shading)
    for r, row in enumerate(rows, start=1):
        cells = t.rows[r].cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10.5)
            run.font.name = "Calibri"
    if col_widths_inches:
        for i, w in enumerate(col_widths_inches):
            for row in t.rows:
                row.cells[i].width = Inches(w)


# ============================================================================
doc = Document()
styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)

# ---- COVER ----
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("MANUALE GENERALE")
r.bold = True
r.font.size = Pt(30)
r.font.color.rgb = ORANGE
r.font.name = "Calibri"

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("App Silvestre Fotoservizi")
r.font.size = Pt(18)
r.font.color.rgb = DARK
r.font.name = "Calibri"

st = doc.add_paragraph()
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st.add_run("La guida che apri quando ti serve capire qualcosa")
r.font.size = Pt(12)
r.italic = True
r.font.color.rgb = GREY
r.font.name = "Calibri"

dt = doc.add_paragraph()
dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = dt.add_run(f"v1.0 — {date.today().strftime('%d/%m/%Y')}")
r.font.size = Pt(10)
r.font.color.rgb = GREY
r.font.name = "Calibri"

doc.add_paragraph()
p(doc, "Indice rapido", bold=True)
bul(doc, [
    "1. La tua app in 1 minuto",
    "2. Firebase, Cloudinary, Pexels, GitHub spiegati",
    "3. La tua app è già online",
    "4. Cose da fare per il rilascio ufficiale (TIER 1 / 2 / 3)",
    "Appendice A — Conformità Legale e GDPR",
    "Appendice B — Guida Firebase (backend dati)",
    "Appendice C — Pubblicazione App Store / Play Store",
    "Appendice D — Deploy FCM e Cloud Functions",
])

doc.add_page_break()

# ============================================================================
h(doc, "1. La tua app in 1 minuto", 1)

p(doc,
  "Hai una sola app con DUE FACCE: clienti e operatori. Cambia in base a chi si logga.")

p(doc, "Cliente:", bold=True)
p(doc, "Apre app → registra → sceglie prodotto → carica foto → ordina → paga (carta/Satispay/in negozio) → "
       "riceve codice ritiro → quando l'ordine è pronto riceve notifica → ritira in negozio.")

p(doc, "Operatore (tu + 3 dipendenti):", bold=True)
p(doc, "Apre app → vede dashboard con ordini in real-time → lavora gli ordini → cambia stato → "
       "manda messaggio cliente via WhatsApp/SMS/Email → cliente notificato.")

p(doc, "Hai 4 documenti in docs/. QUESTO è l'unico che devi leggere sempre. "
       "Gli altri li apri solo quando servono (vedi sezione 7).", italic=True)

doc.add_page_break()

# ============================================================================
h(doc, "2. Firebase, Cloudinary, Pexels e GitHub — chi sono e perché", 1)

p(doc, "La tua app è solo l'INTERFACCIA. I dati e le foto vivono altrove, "
       "in due servizi cloud che ti affittano spazio e funzioni. Senza di loro, "
       "ogni installazione dell'app sarebbe isolata, senza memoria.")

p(doc, "Servizi esterni in cui SEI GIÀ registrato:", bold=True)
table(doc,
      ["Servizio", "Tipo registrazione", "Costo attuale", "A cosa serve"],
      [
          ["Firebase (Google)", "Account Google", "Blaze (pay-as-you-go) ~0-5\u20ac/mese", "Auth + database + push + hosting + Cloud Functions"],
          ["Cloudinary", "Email + password", "Gratis (Developer)", "Magazzino foto utenti"],
          ["Pexels", "Email + password", "Gratis", "Foto stock per catalogo (con attribuzione automatica autori)"],
          ["GitHub", "Account GitHub", "Gratis", "Backup codice + CI/CD auto-deploy"],
      ])

doc.add_paragraph()
h(doc, "\U0001f525 Firebase (Google)", 2)
p(doc, "Il \"cervello\" centrale dell'app. Fa 3 lavori:", bold=True)
bul(doc, [
    "Auth: gestisce login/registrazione. Le password sono cifrate da Google, non da te.",
    "Firestore: il database. Ordini, profili, ruoli, impostazioni — tutto lì.",
    "Cloud Functions: piccoli script automatici. Sono FONDAMENTALI per: push notification reali (FCM) ai clienti app, cron job di pulizia automatica dei pending soft opt-in dopo 30 giorni, eventuali webhook (es. WhatsApp Cloud API se in futuro attivi l'invio automatizzato). Oggi non attivi: serve passare al piano Blaze. Vedi sezione 4 (TIER 2) e Appendice D per il setup.",
])
p(doc, "Perché è importante:", bold=True)
bul(doc, [
    "Cliente e operatore vedono gli STESSI dati in real-time (sync ~1 secondo).",
    "Le password non le gestisci tu = nessun rischio data breach lato tuo.",
    "Quando il cliente torna in app dopo 6 mesi, ritrova i suoi ordini.",
])
p(doc, "Costo: dal 20 maggio 2026 il progetto e' configurato col piano BLAZE "
       "(pay-as-you-go). Significa: stesse soglie gratuite del piano Spark "
       "(50k reads/giorno, 20k writes/giorno, 1 GiB storage, 2M Cloud "
       "Functions invocazioni/mese, FCM ILLIMITATO), ma niente piu' errori "
       "429 quando si supera la soglia istantanea. Si paga solo sopra "
       "soglia: per Silvestre realisticamente 0\u20ac/mese, max 1-3\u20ac/mese "
       "se l'app cresce. Il vero motivo dell'upgrade non e' il costo "
       "(e' 0\u20ac) ma sbloccare Cloud Functions e rate limits adeguati.", bold=True)

doc.add_paragraph()
h(doc, "\u2601\ufe0f Cloudinary", 2)
p(doc, "Il \"magazzino delle foto\". È un servizio specializzato per immagini.", bold=True)
bul(doc, [
    "Quando il cliente carica una foto da stampare, va su Cloudinary, non sul tuo server.",
    "Cloudinary fa anche thumbnail e ottimizzazioni al volo (lo stesso URL serve foto di dimensioni diverse).",
    "Il file è sicuro: solo il cliente e l'operatore possono accedervi.",
])

doc.add_paragraph()
h(doc, "\U0001f4f7 Pexels", 2)
p(doc, "Le foto del catalogo (immagini accanto a ogni prodotto e categoria).", bold=True)
bul(doc, [
    "Pexels è una piattaforma di foto stock con licenza gratuita per uso commerciale.",
    "36 immagini scaricate via API e cachate su Firestore (collection catalog_images).",
    "Ogni scheda prodotto mostra l'attribuzione cliccabile 'Photo by [Autore] - Pexels'.",
    "L'app forza il fetch dal server (Source.server) per vedere subito gli aggiornamenti.",
])
p(doc, "Come si rinfrescano:", bold=True)
bul(doc, [
    "GitHub Actions ha un workflow manuale 'Refresh Catalog Images (Pexels)'.",
    "Triggera dal browser su https://github.com/nico85-prog/silvestre-app/actions",
    "Le queries di ricerca sono in docs/fetch_pexels_images.py (modificabili).",
])
p(doc, "Quando avrai tue foto reali dei prodotti, basta sostituire l'URL nei docs Firestore "
       "catalog_images: l'attribuzione Pexels sparisce automaticamente se il campo "
       "photographer è vuoto. Costo: GRATIS, niente carta richiesta.")

doc.add_paragraph()
h(doc, "\U0001f5c2\ufe0f GitHub", 2)
p(doc, "Dove vive il codice + l'automazione che pubblica gli aggiornamenti.", bold=True)
bul(doc, [
    "Repo privato: nico85-prog/silvestre-app (codice, configurazioni, listini).",
    "GitHub Actions = CI/CD: ad ogni push del codice, il sistema esegue build + test + deploy "
        "automatico su Firebase Hosting. Tu non devi fare nulla.",
    "Storico commit: ogni modifica ha una traccia recuperabile (utile per rollback).",
    "Workflow secondari: refresh immagini catalogo (Pexels) triggerabile manualmente.",
])
p(doc, "Cosa contiene la repo:", bold=True)
bul(doc, [
    "app/ — codice Flutter (cliente + operatore in un'unica app)",
    "firebase/ — config + security rules + functions",
    "docs/ — manuali (questo .docx) + script di seed",
    "Listini/ — CSV catalogo prezzi (modificabile in Excel)",
    ".github/workflows/ — automazione deploy + refresh immagini",
])
p(doc, "Costo: GRATIS per repository privati (piani Free) e CI/CD fino a 2000 minuti/mese "
       "(usiamo circa 5 minuti per deploy).")

doc.add_page_break()

# ============================================================================
h(doc, "3. La tua app è GIÀ ONLINE", 1)
p(doc, "L'app è pubblicata sul cloud di Google ed è raggiungibile da qualsiasi telefono o computer "
       "del mondo. Niente Apple/Google store ancora — solo URL web.")

doc.add_paragraph()
p(doc, "URL pubblico:", bold=True)
p(doc, "  https://silvestre-fotoservizi.web.app", bold=True, color=BLUE)

doc.add_paragraph()
p(doc, "Come installarla \"come app\" su iPhone:", bold=True)
num(doc, [
    "Apri Safari (NON Chrome — Apple permette \"Aggiungi a Home\" solo da Safari)",
    "Vai a https://silvestre-fotoservizi.web.app",
    "Tocca icona Condividi (quadrato con freccia su, in basso al centro)",
    "Scorri menu \u2192 \"Aggiungi alla schermata Home\"",
    "Conferma \"Silvestre\" \u2192 Aggiungi",
    "Esci da Safari \u2192 trovi l'icona arancione sulla home del telefono",
    "Tocca \u2192 si apre a tutto schermo come un'app vera",
])

p(doc, "Come installarla su Android:", bold=True)
num(doc, [
    "Apri Chrome \u2192 vai all'URL sopra",
    "Chrome propone automaticamente \"Installa Silvestre Fotoservizi\" (banner in basso)",
    "Tocca \"Installa\" \u2192 fatto. In alternativa: men\u00f9 3 puntini \u22ee \u2192 \"Installa app\"",
])

doc.add_page_break()

# ============================================================================
h(doc, "4. Cose da fare per il rilascio ufficiale (gratis + a pagamento)", 1)
p(doc, "Tier 1 = obbligatorio prima di accettare ordini reali. "
       "Tier 2 = entro 30 giorni dal lancio. Tier 3 = espansione futura.", italic=True)

doc.add_paragraph()
table(doc,
      ["T", "Cosa", "A cosa serve", "Costo", "Tempo"],
      [
          ["1", "Foto reali dei prodotti", "Le foto attuali sono stock da Pexels: mostrano oggetti generici, non i tuoi reali. I clienti devono riconoscere il prodotto che ricevono, altrimenti il tasso di reso/lamentela sale.", "GRATIS", "1 gg"],
          ["1", "Verifica descrizioni e prezzi catalogo", "Eventuali errori di battitura o prezzi sbagliati ti fanno perdere soldi o credibilità. Una revisione finale prima del lancio è obbligatoria.", "GRATIS", "1 h"],
          ["1", "WhatsApp Business sul negozio", "Ad ogni cambio stato dell'ordine, l'app apre automaticamente WhatsApp con il numero del cliente e il messaggio pre-compilato — l'operatore preme solo Invia. Per professionalità conviene attivare WhatsApp Business sul numero del negozio invece del personale.", "GRATIS", "15 min"],
          ["1", "Bonifico Istantaneo configurato", "Pagamento online a 0% commissioni: IBAN del negozio hardcoded in app (Banca Fideuram, Antonio Silvestre). Cliente paga dalla sua app bancaria, carica ricevuta, operatore conferma in app.", "GRATIS (0% fee)", "Fatto"],
          ["1", "Iubenda Privacy + Termini + Cookie GDPR", "Per legge, app che raccolgono email e dati personali DEVONO avere Privacy Policy e Cookie banner conformi GDPR.", "29-80\u20ac/anno", "30 min"],
          ["1", "Piano Blaze attivato (FATTO)", "Pay-as-you-go senza limiti operativi. Sblocca: Cloud Functions, FCM push reali, cron job automatici, niente piu' errori 429. Costo realistico per Silvestre: 0\u20ac/mese (sotto le soglie gratuite incluse). Soglie free: 50k reads/g, 20k writes/g, 1 GiB storage, 2M Cloud Functions invocazioni/mese.", "~0\u20ac/mese stimato", "5 min (carta)"],
          ["2", "Cloud Functions deploy (FCM + cron)", "Una volta attivo Blaze, deploy delle Cloud Functions per: (a) invio FCM push ai clienti app quando l'operatore crea promo, (b) cron job giornaliero che marca optInStatus=no i pending oltre 30 giorni, (c) eventuale webhook WhatsApp Cloud API. Tutto serverless, nessuna manutenzione.", "Inclusi nel Blaze", "2-3 h sviluppo"],
          ["2", "Aruba/FattureInCloud", "La normativa italiana richiede fatturazione elettronica oltre soglia.", "25-50\u20ac/anno", "1 h"],
          ["3", "Apple Developer", "Per pubblicare sull'App Store iPhone.", "99$/anno", "2 sett."],
          ["3", "Google Play Developer", "Per pubblicare sul Play Store Android.", "25$ una tantum", "3 sett."],
          ["3", "Build + submission iOS/Android", "Convertire la PWA in app native, generare screenshot, descrizioni store.", "GRATIS (dopo Apple/Google)", "3-5 gg"],
          ["3", "Polizze RC + Cyber + GDPR audit", "Protezione legale per dispute e violazioni dati. Audit annuale per restare conforme.", "700-1500\u20ac/anno", "1 h"],
      ],
      # T: stretta (0.35"), Cosa: media (1.5"), A cosa serve: larga (3.4"), Costo: 1.2", Tempo: 0.8"
      col_widths_inches=[0.35, 1.5, 3.4, 1.2, 0.8])

doc.add_paragraph()
p(doc, "STIMA COSTI ANNO 1", bold=True)
table(doc,
      ["Scenario", "Totale annuo"],
      [
          ["Minimo per partire (Iubenda + dominio)", "~90\u20ac"],
          ["Professional base (+ store + email + fatture)", "~400\u20ac"],
          ["Professional full (+ assicurazioni)", "~1.700\u20ac"],
      ])

doc.add_paragraph()
p(doc, "GIÀ A POSTO (non devi fare nulla)", bold=True)
bul(doc, [
    "Auth Firebase + telefono OBBLIGATORIO + un solo account operatore condiviso (operatore@silvestrefotoservizi.it) con permessi pieni",
    "Catalogo 30 prodotti / 193 varianti con tutte le fasce qty (modificabile via CSV)",
    "Carrello editabile + ordini real-time + photobook auto-impagina AI + lavoro personalizzato",
    "Auto-WhatsApp al cliente ad ogni cambio stato ordine (operatore preme solo Invia)",
    "Bottone WhatsApp + mappa Google integrati per contatto rapido cliente",
    "Operatore: dashboard, ordini, preventivi, gestione operatori, template messaggi",
    "Pannello operatore con LINK DIRETTI a: Google Analytics, Firebase Console, Cloudinary",
    "Multilingua IT/EN con switch lingua in Account (UI tradotta, catalogo resta IT)",
    "PWA installabile iOS/Android + 5 temi personalizzabili",
    "CI/CD GitHub Actions auto-deploy + workflow refresh immagini Pexels",
    "Backup automatico: tutti i dati persistono ai deploy (mai resettati)",
    "Sentry-ready: serve solo creare account gratis + incollare DSN",
    "Google Analytics-ready: serve solo creare account + incollare measurement ID",
])

doc.add_paragraph()
p(doc, "PROSSIMO PASSO CONSIGLIATO", bold=True)
bul(doc, [
    "Foto reali dei prodotti (gratis, risultato visibile subito)",
    "Iubenda (30 minuti, sblocchi tutto il GDPR)",
    "Beta test con amici (manda Istruzioni_Amici_Test.txt)",
])

doc.add_paragraph()
p(doc, "Fine manuale. Per qualsiasi cosa tecnica scrivi a me. Per cose legali/fiscali al "
       "tuo avvocato/commercialista.", italic=True)
p(doc, f"Silvestre Fotoservizi \u00b7 {date.today().strftime('%d/%m/%Y')} \u00b7 v1.0",
  italic=True, color=GREY, size=10)

import os
out = os.path.join(os.path.dirname(__file__), "1_Manuale_Generale_Silvestre.docx")
doc.save(out)
print(f"OK: {out}")
