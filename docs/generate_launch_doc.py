"""
Genera il documento Word "Pubblicazione_App_Silvestre.docx" con guida
passo-passo per il rilascio dell'app su App Store (iOS) e Google Play (Android).

ATTENZIONE: documento informativo. Le procedure di Apple/Google possono cambiare;
verificare sempre le linee guida ufficiali al momento del rilascio.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date

ORANGE = RGBColor(0xF4, 0x75, 0x21)
DARK = RGBColor(0x2B, 0x2B, 0x2B)
GREY = RGBColor(0x7A, 0x7A, 0x7A)
RED = RGBColor(0xD6, 0x45, 0x45)
GREEN = RGBColor(0x3D, 0xA3, 0x5D)


def add_heading(doc, text, level=1, color=ORANGE):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = "Calibri"


def add_p(doc, text, bold=False, italic=False, color=DARK, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    return p


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(it)
        run.font.size = Pt(11)
        run.font.name = "Calibri"


def add_numbered(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(it)
        run.font.size = Pt(11)
        run.font.name = "Calibri"


def add_checklist(doc, items):
    for it in items:
        p = doc.add_paragraph()
        chk = p.add_run("☐  ")
        chk.font.size = Pt(13)
        chk.bold = True
        chk.font.color.rgb = ORANGE
        txt = p.add_run(it)
        txt.font.size = Pt(11)
        txt.font.name = "Calibri"


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F47521")
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
    for r, row in enumerate(rows, start=1):
        cells = t.rows[r].cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10.5)
            run.font.name = "Calibri"


def add_warning(doc, text):
    p = doc.add_paragraph()
    run = p.add_run("⚠  " + text)
    run.bold = True
    run.font.color.rgb = RED
    run.font.size = Pt(11)
    run.font.name = "Calibri"


def add_tip(doc, text):
    p = doc.add_paragraph()
    run = p.add_run("💡  " + text)
    run.font.color.rgb = GREEN
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Calibri"


def add_spacer(doc):
    doc.add_paragraph("")


# ============================================================================
doc = Document()
styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

# --- COPERTINA ---
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Pubblicazione App Mobile")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = ORANGE
r.font.name = "Calibri"

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Guida pratica per Android (Google Play) e iOS (App Store)")
r.font.size = Pt(15)
r.font.color.rgb = DARK
r.font.name = "Calibri"

brand = doc.add_paragraph()
brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = brand.add_run("Silvestre Fotoservizi — Frattamaggiore (NA)")
r.font.size = Pt(13)
r.italic = True
r.font.color.rgb = GREY
r.font.name = "Calibri"

dt = doc.add_paragraph()
dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = dt.add_run(f"Versione del {date.today().strftime('%d/%m/%Y')}")
r.font.size = Pt(10)
r.font.color.rgb = GREY
r.font.name = "Calibri"

add_spacer(doc)

add_p(
    doc,
    "Questa guida copre i 3 modi per pubblicare l'app Silvestre Fotoservizi: "
    "PWA web (gratis), App Store Apple (a pagamento), Play Store Google (a pagamento). "
    "Va letta UNA volta all'inizio e riconsultata prima di ogni rilascio.",
    italic=True,
)
add_spacer(doc)
add_tip(
    doc,
    "STRATEGIA CONSIGLIATA: parti con la PWA web (sez. 1.5 sotto) — costa solo il dominio "
    "(~15 EUR/anno) e ti permette di lanciare in 1 ora. Aggiungi gli store mobile dopo, "
    "quando i clienti lo chiedono."
)

doc.add_page_break()

# --- PWA WEB (NEW SECTION) ---
add_heading(doc, "1.5 OPZIONE A — PWA web (la più economica)", level=1)
add_p(
    doc,
    "Una PWA (Progressive Web App) è una applicazione che vive su un sito web ma "
    "puo' essere INSTALLATA come app sul telefono. Su iPhone tramite menu Condividi "
    "-> Aggiungi a Home; su Android Chrome propone automaticamente Installa.",
)

add_heading(doc, "Cosa serve", level=2)
add_table(
    doc,
    ["Voce", "Costo", "Dove"],
    [
        ["Dominio .it (es. app.silvestrefotoservizi.it)", "10-20 EUR/anno", "aruba.it, register.it, dynadot.com"],
        ["Hosting Firebase", "0 EUR (free tier)", "Gia' incluso nel tuo progetto Firebase"],
        ["Certificato SSL", "0 EUR", "Automatico via Firebase Hosting"],
        ["TOTALE", "~15 EUR/anno", ""],
    ],
)

add_heading(doc, "Stato attuale", level=2)
add_p(doc, "L'app E' GIA' DEPLOYATA su Firebase Hosting:")
add_p(doc, "  https://silvestre-fotoservizi.web.app", bold=True)
add_p(doc, "L'unico \"upgrade\" futuro e' aggiungere un dominio personalizzato (es. "
           "app.silvestrefotoservizi.it) tramite Firebase Console -> Hosting -> "
           "\"Add custom domain\". Costa solo il dominio (~15 EUR/anno).")

add_heading(doc, "Come pubblicare la PWA (passi tecnici)", level=2)
add_numbered(
    doc,
    [
        "Compra dominio su Aruba/Register",
        "Da firebase.cmd: 'firebase init hosting' nella cartella firebase/",
        "Configura public directory: ../app/build/web",
        "Build: cd app && flutter build web --release",
        "Deploy: cd ../firebase && firebase deploy --only hosting",
        "Su Firebase Console -> Hosting -> Add custom domain -> inserisci il tuo dominio",
        "Aggiorna i record DNS del dominio come da istruzioni Firebase",
        "Attendi 24-48h propagazione DNS",
        "L'app e' online su https://app.silvestrefotoservizi.it",
    ],
)
add_tip(
    doc,
    "Tutto questo lo puo' fare il tuo tecnico (o io quando glielo chiedi) in 1-2 ore. "
    "Niente Apple Developer, niente Google Play, niente attese."
)

add_warning(
    doc,
    "Limiti PWA su iOS: niente push notification (iOS le supporta solo per app store nativa). "
    "Su Android invece le push funzionano anche con PWA. Per le push iOS DEVI passare alla "
    "Strada B (App Store)."
)

doc.add_page_break()

# --- 1. PANORAMICA COSTI ---
add_heading(doc, "1. OPZIONE B — Store nativi (Apple + Google) — Costi", level=1)

add_table(
    doc,
    ["Voce", "Costo", "Frequenza"],
    [
        ["Apple Developer Program", "99 € / 109 USD", "Annuale (rinnovo automatico)"],
        ["Google Play Developer Account", "25 USD (~23 €)", "Una sola volta, mai più"],
        ["Mac (richiesto per build iOS)", "0 €", "Già posseduto"],
        ["Servizio cloud build (alternativa Mac)", "30-50 €/mese", "Se non usi il tuo Mac"],
        ["Strumenti di sviluppo (Xcode, Android Studio)", "0 €", "Gratuiti"],
        ["Dominio per Privacy Policy URL", "10-20 €/anno", "Annuale (es. silvestrefotoservizi.it se non hai già)"],
        ["Certificati SSL", "0 €", "Inclusi nel dominio Cloudflare o gratis con Let's Encrypt"],
        ["Servizio di firma app (Fastlane, opzionale)", "0 €", "Open source"],
    ],
)
add_spacer(doc)
add_p(
    doc,
    "TOTALE primo anno: ~130 € (Apple + Google + dominio). Anno 2+: ~110 €/anno "
    "(solo Apple + dominio). Google è una tantum.",
    bold=True,
)

doc.add_page_break()

# --- 2. ACCOUNT SVILUPPATORE APPLE ---
add_heading(doc, "2. Account sviluppatore Apple (iOS)", level=1)

add_heading(doc, "2.1 Cosa devi avere", level=2)
add_bullets(
    doc,
    [
        "Apple ID personale (gratis, se non ce l'hai crealo su appleid.apple.com).",
        "Un Mac per fare le build (sì, obbligatorio per Apple; tu hai detto di averlo).",
        "Codice fiscale o partita IVA (se ti registri come azienda).",
        "Carta di credito per il pagamento annuale.",
        "DUNS Number se ti registri come azienda anziché come individuo "
        "(gratis, va richiesto su dnb.com — ci vogliono 2 settimane). "
        "Se ti registri come individuo, NON serve.",
    ],
)

add_heading(doc, "2.2 Individuale o Azienda?", level=2)
add_table(
    doc,
    ["Tipo", "Pro", "Contro"],
    [
        [
            "Individuale (Individual)",
            "Iscrizione veloce (1-2 giorni), no DUNS",
            "Nello store appare il tuo NOME, non Silvestre Fotoservizi",
        ],
        [
            "Azienda (Organization)",
            "Nello store appare 'Silvestre Fotoservizi', più professionale",
            "Serve DUNS (2 settimane), più documenti",
        ],
    ],
)
add_tip(
    doc,
    "Per il negozio Silvestre Fotoservizi conviene iscriversi come AZIENDA "
    "(il nome viene visto da tutti i clienti che scaricano l'app).",
)

add_heading(doc, "2.3 Passi per la registrazione", level=2)
add_numbered(
    doc,
    [
        "Vai su developer.apple.com/programs/enroll/",
        "Accedi con il tuo Apple ID.",
        "Se Azienda: inserisci dati legali (P.IVA, ragione sociale Silvestre Fotoservizi S.a.s., DUNS).",
        "Conferma i termini.",
        "Paga 99 €/anno con carta di credito.",
        "Attesa approvazione: 24-48h per Individual, 5-15 giorni per Organization.",
        "Ricevi email di conferma → accedi a App Store Connect (appstoreconnect.apple.com).",
    ],
)

doc.add_page_break()

# --- 3. ACCOUNT SVILUPPATORE GOOGLE ---
add_heading(doc, "3. Account sviluppatore Google Play (Android)", level=1)

add_heading(doc, "3.1 Cosa devi avere", level=2)
add_bullets(
    doc,
    [
        "Account Google (puoi usare lo stesso del progetto Firebase).",
        "Carta di credito per il pagamento una tantum.",
        "Documento di identità (per verifica).",
        "Indirizzo aziendale per i clienti (verrà mostrato nel listing).",
    ],
)

add_heading(doc, "3.2 Passi per la registrazione", level=2)
add_numbered(
    doc,
    [
        "Vai su play.google.com/console/signup",
        "Accedi con account Google.",
        "Scegli tipo account: Personale o Aziendale (consigliato Aziendale per Silvestre).",
        "Per Aziendale: inserisci dati (Silvestre Fotoservizi S.a.s., P.IVA).",
        "Paga 25 USD con carta.",
        "Verifica identità: carica documento (carta d'identità o passaporto).",
        "Verifica indirizzo aziendale (Google invia una cartolina con codice in 2-4 settimane).",
        "Una volta verificato → Google Play Console pronta.",
    ],
)
add_warning(
    doc,
    "Dal 2024 Google richiede verifica per cartolina. Senza, NON puoi pubblicare app a pagamento "
    "né app che gestiscono pagamenti. Inizia la verifica SUBITO appena registrato.",
)

doc.add_page_break()

# --- 4. MATERIALI ---
add_heading(doc, "4. Materiali da preparare (per entrambi gli store)", level=1)

add_heading(doc, "4.1 Icona app", level=2)
add_table(
    doc,
    ["Asset", "Misura", "Formato"],
    [
        ["Icona App Store", "1024x1024 px", "PNG senza trasparenza"],
        ["Icona Google Play", "512x512 px", "PNG 32-bit con alfa"],
        ["Icona app (varianti)", "Generate da Flutter", "Generate da flutter_launcher_icons"],
        ["Adaptive icon Android (foreground)", "432x432 px", "PNG con alfa"],
        ["Adaptive icon Android (background)", "432x432 px", "Colore o PNG"],
    ],
)
add_tip(
    doc,
    "Useremo il pacchetto Flutter 'flutter_launcher_icons' per generare automaticamente "
    "tutte le varianti partendo dal tuo logo silvestre_logo.jpg.",
)

add_heading(doc, "4.2 Screenshots", level=2)
add_p(doc, "Servono screenshot per le dimensioni dei dispositivi venduti:")
add_table(
    doc,
    ["Store", "Dispositivo", "Risoluzione", "Quantità minima"],
    [
        ["App Store", "iPhone 6.7\" (15 Pro Max)", "1290x2796 px", "3 (max 10)"],
        ["App Store", "iPhone 6.5\" (XS Max)", "1242x2688 px", "3"],
        ["App Store", "iPad 13\" (opzionale)", "2064x2752 px", "Solo se supporti iPad"],
        ["Google Play", "Telefono", "1080x1920 o sup.", "2 (max 8)"],
        ["Google Play", "Tablet 7\" (opz.)", "Variabile", "Solo se ottimizzi tablet"],
        ["Google Play", "Tablet 10\" (opz.)", "Variabile", "Solo se ottimizzi tablet"],
        ["Google Play", "Feature Graphic", "1024x500 px", "1 obbligatoria"],
    ],
)
add_tip(
    doc,
    "Possiamo generare gli screenshot automaticamente dal simulatore con uno script Flutter, "
    "oppure usare servizi come AppScreens.io o ScreenshotOne.",
)

add_heading(doc, "4.3 Testi (in italiano)", level=2)
add_bullets(
    doc,
    [
        "Nome app (30 caratteri Apple, 30 Google): es. 'Silvestre Fotoservizi'.",
        "Sottotitolo (30 caratteri Apple): es. 'Stampe a Frattamaggiore dal 1970'.",
        "Promozionale breve (170 caratteri Apple, 80 Google): es. 'Ordina stampe, fotolibri, calendari. Ritiro in negozio.'",
        "Descrizione lunga (4000 caratteri): cosa fa l'app, prodotti, come funziona, info negozio, FAQ.",
        "Keywords iOS (100 caratteri totali, separati da virgola): 'foto,stampa,fotolibro,calendario,Frattamaggiore,Napoli'.",
        "Categoria: Foto e Video.",
        "URL Privacy Policy obbligatorio (deve essere pubblico e raggiungibile).",
        "URL supporto obbligatorio (email o pagina contatti).",
        "URL marketing (opzionale): il tuo sito silvestrefotoservizi.it.",
    ],
)

add_heading(doc, "4.4 Classificazione contenuti", level=2)
add_p(doc, "Compili un questionario per età:")
add_bullets(
    doc,
    [
        "Apple: età suggerita 4+ (nessun contenuto inappropriato).",
        "Google: completa IARC questionnaire (5 minuti).",
        "Se l'app permette upload di contenuti utente devi dichiarare 'User Generated Content'.",
    ],
)

doc.add_page_break()

# --- 5. PRIVACY LABELS ---
add_heading(doc, "5. Etichette privacy (obbligatorie e RIGOROSE)", level=1)

add_warning(
    doc,
    "Mentire o omettere nelle privacy labels = rejection o ban. Apple e Google fanno controlli automatici.",
)

add_heading(doc, "5.1 Apple — Privacy Nutrition Labels", level=2)
add_p(
    doc,
    "In App Store Connect, sezione 'App Privacy', dichiari per ogni tipo di dato: "
    "Raccolto / Non raccolto, e per ognuno raccolto: finalità + se è collegato all'utente "
    "+ se è usato per tracking.",
)
add_p(doc, "Per la nostra app dichiariamo:")
add_table(
    doc,
    ["Dato raccolto", "Finalità", "Collegato all'utente", "Tracking"],
    [
        ["Email", "Funzionalità app", "Sì", "No"],
        ["Nome", "Funzionalità app", "Sì", "No"],
        ["Telefono", "Funzionalità app", "Sì", "No"],
        ["Foto e video", "Funzionalità app", "Sì", "No"],
        ["Cronologia acquisti", "Funzionalità app, Analytics", "Sì", "No"],
        ["Identificatori dispositivo", "Analytics", "Sì", "No"],
    ],
)

add_heading(doc, "5.2 Google — Data Safety", level=2)
add_p(
    doc,
    "In Google Play Console, sezione 'Data safety', stesso principio. Più "
    "dichiari se i dati sono criptati in transit (HTTPS sì, Firebase) e se "
    "l'utente può richiedere la cancellazione (sì, abbiamo implementato Elimina account).",
)

doc.add_page_break()

# --- 6. BUILD APP ---
add_heading(doc, "6. Generare la build di release", level=1)

add_heading(doc, "6.1 Android (APK / AAB)", level=2)
add_numbered(
    doc,
    [
        "Genera una keystore di firma (una sola volta nella vita dell'app, conservarla GELOSAMENTE).",
        "Configura android/key.properties con la keystore.",
        "Esegui: flutter build appbundle --release",
        "Output: build/app/outputs/bundle/release/app-release.aab",
        "Carica il .aab in Google Play Console.",
    ],
)
add_warning(
    doc,
    "Se PERDI la keystore Android, NON puoi più aggiornare l'app. Conservala in 3 posti diversi "
    "(disco, cloud criptato, USB in cassetta sicurezza). Salva anche la password.",
)

add_heading(doc, "6.2 iOS (IPA)", level=2)
add_numbered(
    doc,
    [
        "Sul Mac, apri ios/Runner.xcworkspace in Xcode.",
        "In Xcode: Signing & Capabilities → seleziona il tuo Team (Silvestre Fotoservizi).",
        "Esegui: flutter build ios --release",
        "In Xcode: Product → Archive → Distribute App → App Store Connect.",
        "L'IPA viene caricato automaticamente su App Store Connect.",
    ],
)
add_tip(
    doc,
    "Possiamo automatizzare entrambe le build con Fastlane (open source) o GitHub Actions, "
    "così basta un comando e fa tutto: build + upload + tag git.",
)

doc.add_page_break()

# --- 7. INVIO REVIEW ---
add_heading(doc, "7. Invio per revisione (review)", level=1)

add_heading(doc, "7.1 Apple", level=2)
add_p(
    doc,
    "Dopo aver caricato la build e compilato tutti i metadati, premi 'Add for Review'. "
    "Apple esamina l'app entro 24-72 ore in media (a volte fino a 7 giorni). "
    "Se rifiutata, ti danno motivazione → correggi → rinvii (gratis, illimitato).",
)

add_p(doc, "Motivi più comuni di rejection nella nostra categoria:")
add_bullets(
    doc,
    [
        "Privacy Policy URL non raggiungibile.",
        "Privacy Labels incomplete o inaccurate.",
        "App che usa fotocamera senza spiegare perché in info.plist (NSCameraUsageDescription).",
        "App che permette upload utente senza sistema di segnalazione abusi (1.2 guideline).",
        "Crash all'apertura su qualche dispositivo testato dal reviewer.",
        "Login obbligatorio ma niente account demo per il reviewer (richiesto in App Review Notes).",
    ],
)
add_tip(
    doc,
    "Crea SEMPRE un account demo (es. demo@silvestre.it / DemoSilv2026!) e mettilo nelle 'App Review Notes' "
    "di App Store Connect. Il reviewer lo userà.",
)

add_heading(doc, "7.2 Google Play", level=2)
add_p(
    doc,
    "Google ha un percorso più graduale tramite tracks:",
)
add_numbered(
    doc,
    [
        "Internal testing: rilascio in pochi minuti, max 100 tester (lista di email).",
        "Closed testing: tester più ampio, review automatica (qualche ora).",
        "Open testing: chiunque può scaricare via link beta, review più lunga (1-7 giorni).",
        "Production: rilascio pubblico, review completa (1-3 giorni in media, "
        "ma può richiedere 14+ giorni la prima volta).",
    ],
)
add_warning(
    doc,
    "Dal 2024 Google impone Closed Testing con almeno 12 tester per almeno 14 giorni "
    "PRIMA del rilascio in produzione per nuovi account. Pianifica con anticipo.",
)

doc.add_page_break()

# --- 8. POST-PUBBLICAZIONE ---
add_heading(doc, "8. Dopo la pubblicazione", level=1)

add_heading(doc, "8.1 Monitoraggio", level=2)
add_bullets(
    doc,
    [
        "App Store Connect → Analytics: download, installazioni, crashes.",
        "Google Play Console → Statistics + Vitals: ANR, crash rate, performance.",
        "Firebase Crashlytics (gratis): stack trace di ogni crash con frequenza.",
        "Sentry (opz, gratis sotto 5k eventi/mese): errori non bloccanti.",
    ],
)

add_heading(doc, "8.2 Recensioni", level=2)
add_bullets(
    doc,
    [
        "Rispondi sempre alle recensioni negative entro 48h.",
        "Rispondi anche a quelle positive per fidelizzare.",
        "Le recensioni pubbliche influenzano la posizione nelle ricerche.",
    ],
)

add_heading(doc, "8.3 Aggiornamenti", level=2)
add_bullets(
    doc,
    [
        "Ad ogni nuova versione: aumenta version + buildNumber in pubspec.yaml.",
        "Build nuova → carica → metti What's New (note di rilascio in italiano).",
        "Submit per review (stessa procedura).",
        "Aggiornamenti minori passano review più velocemente (~12-24h Apple).",
        "Aggiornamenti che cambiano trattamento dati richiedono aggiornare Privacy Labels.",
    ],
)

doc.add_page_break()

# --- 9. CHECKLIST PRE-RILASCIO ---
add_heading(doc, "9. Checklist completa pre-rilascio", level=1)

add_heading(doc, "Account e legal", level=2)
add_checklist(
    doc,
    [
        "Apple Developer Account attivo (99€ pagati).",
        "Google Play Developer Account attivo (25€ pagati + verifica completata).",
        "P.IVA inserita correttamente nei due account.",
        "Privacy Policy pubblicata a URL pubblica.",
        "Termini di Servizio pubblicati a URL pubblica.",
        "Email di supporto attiva (supporto@silvestrefotoservizi.it).",
        "DPA firmati con Firebase (automatico) e payment provider.",
    ],
)

add_heading(doc, "Materiali grafici", level=2)
add_checklist(
    doc,
    [
        "Icona 1024x1024 (Apple).",
        "Icona 512x512 (Google).",
        "Adaptive icon Android (foreground + background).",
        "Screenshot iPhone 6.7\" (3+).",
        "Screenshot iPhone 6.5\" (3+).",
        "Screenshot Android phone (2+).",
        "Feature Graphic Google 1024x500.",
    ],
)

add_heading(doc, "Testi", level=2)
add_checklist(
    doc,
    [
        "Nome app deciso.",
        "Sottotitolo Apple deciso (max 30 char).",
        "Promozionale breve scritto.",
        "Descrizione lunga scritta (in italiano).",
        "Keywords iOS scelte.",
        "Note di rilascio v1.0.0 scritte.",
    ],
)

add_heading(doc, "Tecnici", level=2)
add_checklist(
    doc,
    [
        "Bundle ID iOS: com.silvestrefotoservizi.app.",
        "Application ID Android: com.silvestrefotoservizi.app.",
        "Version: 1.0.0+1 in pubspec.yaml.",
        "Icone generate con flutter_launcher_icons.",
        "Splash screen configurato (flutter_native_splash).",
        "Permessi richiesti dichiarati: NSCameraUsageDescription, NSPhotoLibraryUsageDescription (iOS).",
        "Permessi Android: CAMERA, READ_MEDIA_IMAGES.",
        "Build release testata su 3+ dispositivi reali.",
        "Nessun warning in flutter analyze.",
        "Tutti i test passano (flutter test).",
        "Crashlytics o Sentry configurato.",
    ],
)

add_heading(doc, "Privacy labels", level=2)
add_checklist(
    doc,
    [
        "Apple Privacy Labels compilate accuratamente.",
        "Google Data Safety compilato accuratamente.",
        "Lista dati raccolti coerente con Privacy Policy.",
    ],
)

add_heading(doc, "Pre-submission", level=2)
add_checklist(
    doc,
    [
        "Account demo creato e inserito in App Review Notes Apple.",
        "TestFlight beta inviato a 2-3 amici per test (Apple).",
        "Internal testing Google con 5+ persone per 1 settimana.",
        "Tutte le funzioni del flusso utente testate end-to-end.",
        "Backup keystore Android salvato in 3 posti.",
        "Certificati Apple esportati e salvati.",
    ],
)

doc.add_page_break()

# --- 10. TIMELINE REALISTA ---
add_heading(doc, "10. Timeline realista dal codice al primo download", level=1)

add_table(
    doc,
    ["Fase", "Tempo", "Note"],
    [
        ["Registrazione Apple (Individual)", "2-3 giorni", "Veloce"],
        ["Registrazione Apple (Org + DUNS)", "2-3 settimane", "DUNS è il collo di bottiglia"],
        ["Registrazione Google + verifica cartolina", "3-5 settimane", "Cartolina arriva da USA"],
        ["Preparazione materiali (icone, screenshot, testi)", "2-3 giorni", "Si può parallelizzare"],
        ["Privacy Policy + Termini (Iubenda)", "1 giorno", "Compili form online"],
        ["Build di release + test", "1-2 giorni", "Test su dispositivi reali"],
        ["Internal/Closed testing Google", "14 giorni minimi", "Imposti da Google 2024"],
        ["Submission Apple", "1 giorno", "Upload e compila metadata"],
        ["Review Apple prima volta", "1-7 giorni", "Tipicamente 24-48h"],
        ["Review Google Production", "1-14 giorni", "Più lungo per nuovi account"],
    ],
)
add_spacer(doc)
add_p(
    doc,
    "TIMELINE TOTALE realistica per primo rilascio: 4-8 settimane (se parti da zero su entrambi). "
    "Pianifica di iniziare le registrazioni Apple e Google ORA, in parallelo allo sviluppo, "
    "così non sono il collo di bottiglia alla fine.",
    bold=True,
)

doc.add_page_break()

# --- 11. ERRORI DA EVITARE ---
add_heading(doc, "11. Errori comuni da evitare", level=1)

add_bullets(
    doc,
    [
        "Non rinnovare l'Apple Developer Account → l'app sparisce dallo store il giorno della scadenza.",
        "Perdere la keystore Android → mai più aggiornamenti.",
        "Cambiare Bundle ID dopo il rilascio → diventa un'app nuova, perdi gli utenti.",
        "Dichiarare il falso nelle Privacy Labels → ban del developer account.",
        "Includere SDK di tracking senza dichiararlo → rejection automatica.",
        "Codice di test/debug lasciato attivo in release → rejection o issue.",
        "Hardcode di API keys di Firebase nel codice client → leak (Firebase è OK perché le chiavi client sono pubbliche, ma applica security rules!).",
        "Permessi richiesti senza giustificazione (es. location se non serve) → rejection.",
        "Login obbligatorio senza account demo per reviewer → rejection sicura.",
        "Promesse non mantenute nella descrizione (es. 'a breve disponibile X') → rejection per metadata.",
    ],
)

# --- 12. CHIUSURA ---
add_spacer(doc)
add_warning(
    doc,
    "Le policy di Apple e Google cambiano spesso. Prima del rilascio CONSULTA SEMPRE: "
    "developer.apple.com/app-store/review/guidelines/ e play.google.com/about/developer-content-policy/",
)

import os
out = os.path.join(os.path.dirname(__file__), "Pubblicazione_App_Silvestre.docx")
doc.save(out)
print(f"OK: {out}")
