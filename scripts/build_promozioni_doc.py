"""Genera Promozioni.docx con il piano completo per la feature Promozioni."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0xF4, 0x75, 0x21)  # Silvestre orange
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    return p


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, cell in enumerate(row):
            t.rows[r_idx].cells[c_idx].text = cell
    return t


def add_callout(doc, title, body):
    p = doc.add_paragraph()
    r = p.add_run(f"▸ {title}\n")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xF4, 0x75, 0x21)
    r2 = p.add_run(body)
    r2.font.size = Pt(10)
    r2.italic = True
    return p


def main():
    doc = Document()

    # --- TITOLO ---
    title = doc.add_heading("Silvestre Fotoservizi — Piano Promozioni", level=0)
    for r in title.runs:
        r.font.color.rgb = RGBColor(0xF4, 0x75, 0x21)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Dalla A alla Z: consenso cliente, contatti, canali di invio, "
                  "logica operativa, vincoli legali e tecnici.")
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Documento di pianificazione — versione 1.0 — 20 maggio 2026")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    # ============================================================
    # 1. PREMESSA
    # ============================================================
    add_heading(doc, "1. Premessa e obiettivo", level=1)
    add_para(doc,
             "Silvestre Fotoservizi possiede una base di 6000+ contatti clienti "
             "(file Contatti_Clienti.csv) raccolti negli anni dall'attività in "
             "negozio. L'obiettivo è poter inviare promozioni (sconti stagionali, "
             "offerte lampo, novità) a questa base in modo:")
    add_bullet(doc, "Legale (conforme al GDPR e al provvedimento del Garante "
                    "Privacy in materia di marketing diretto)")
    add_bullet(doc, "Economico (zero costi ricorrenti per il negozio)")
    add_bullet(doc, "Sostenibile nel tempo (l'operatore non deve diventare uno "
                    "schiavo della tastiera per inviare messaggi manualmente)")
    add_bullet(doc, "Tracciato (sapere chi ha ricevuto cosa, chi ha risposto "
                    "SI/STOP, chi è in blacklist)")

    # ============================================================
    # 2. IL PROBLEMA DEL CONSENSO (GDPR)
    # ============================================================
    add_heading(doc, "2. Il problema del consenso (GDPR)", level=1)
    add_para(doc, "2.1 — Cosa dice la legge", bold=True, size=12)
    add_para(doc,
             "Il GDPR (Regolamento UE 2016/679) e il Codice Privacy italiano "
             "richiedono per le comunicazioni promozionali via WhatsApp / email / "
             "SMS un consenso che sia:")
    add_bullet(doc, "Esplicito — il cliente deve aver attivamente spuntato una "
                    "casella «accetto di ricevere promozioni»")
    add_bullet(doc, "Specifico — il consenso al marketing è separato dal consenso "
                    "all'uso del servizio")
    add_bullet(doc, "Libero — non puoi rendere obbligatorio il marketing per "
                    "usare l'app (art. 7(4) GDPR: divieto di «consenso forzato»)")
    add_bullet(doc, "Documentato — devi poter dimostrare quando e come l'hai "
                    "ottenuto")
    add_bullet(doc, "Revocabile — il cliente può sempre disiscriversi (opt-out)")

    add_para(doc, "")
    add_para(doc, "2.2 — Lo stato attuale dei 6000+ contatti CSV", bold=True, size=12)
    add_para(doc,
             "I contatti nel CSV sono stati raccolti durante anni di attività in "
             "negozio. Hanno fornito il numero per ricevere conferme ordine, "
             "comunicazioni di servizio, ma NON hanno mai prestato un esplicito "
             "consenso a ricevere comunicazioni promozionali. Mandare loro una "
             "promo diretta è giuridicamente classificabile come spam, con "
             "sanzioni che possono arrivare a 20 milioni di euro o al 4% del "
             "fatturato annuo (art. 83 GDPR).")

    add_para(doc, "")
    add_para(doc, "2.3 — La soluzione adottata: doppio binario", bold=True, size=12)
    add_para(doc,
             "Per i futuri nuovi clienti che si registrano nell'app: checkbox "
             "marketing OPZIONALE, ben evidenziata graficamente per invogliare a "
             "selezionarla, ma non bloccante (l'app funziona uguale anche se la "
             "saltano). Default OFF, come richiede il GDPR.")
    add_para(doc,
             "Per i 6000+ contatti già esistenti nel CSV: campagna di soft "
             "opt-in. Una singola WhatsApp di richiesta consenso, chi risponde "
             "SI viene aggiunto alla lista marketing, gli altri vengono "
             "automaticamente segnati come non-opted-in dopo 30 giorni e "
             "rimossi dalla lista invii.")

    add_callout(doc, "Effetto pratico",
                "Nel giro di 2-3 mesi avremo una lista pulita di clienti opted-in "
                "(stima realistica: 1.500-2.500 contatti sui 6000+ iniziali). "
                "A quella lista possiamo mandare promozioni in totale "
                "tranquillità legale per anni.")

    add_para(doc, "")
    add_para(doc, "2.4 — Regola d'oro: la scelta esplicita vince sempre", bold=True, size=12)
    add_para(doc,
             "Il sistema applica automaticamente una regola fondamentale: se un "
             "cliente si registra nell'app e NON spunta la checkbox marketing, "
             "il suo eventuale stato 'pending' nel CSV viene immediatamente "
             "convertito in 'no' permanente. La sua scelta esplicita nell'app "
             "ha la priorità sullo stato CSV, e il soft opt-in non gli verrà "
             "mai inviato. Questo è dettagliato nella Tab 4 «Logica & GDPR» del "
             "pannello operatore.")

    # ============================================================
    # 3. IL MESSAGGIO DI SOFT OPT-IN
    # ============================================================
    add_heading(doc, "3. Campagna di soft opt-in iniziale", level=1)

    add_callout(doc, "PERIMETRO ESCLUSIVO",
                "La campagna di soft opt-in si applica UNICAMENTE ai contatti "
                "del CSV storico che NON hanno mai scaricato l'app Silvestre "
                "Fotoservizi. Chi è già registrato nell'app ha già espresso "
                "la propria scelta marketing (sì o no) durante la registrazione, "
                "quindi non riceve mai il messaggio di soft opt-in. Vedi "
                "sezione 5.2.2 «Le 4 categorie di cliente» per il dettaglio "
                "delle regole di dedup automatico.")

    add_para(doc, "")
    add_para(doc, "3.1 — Template del messaggio", bold=True, size=12)
    add_para(doc,
             "Il messaggio inviato una singola volta a ciascun contatto del "
             "CSV NON registrato in app:")
    p = doc.add_paragraph()
    r = p.add_run(
        "« Ciao {{nome}}, ti scriviamo da Silvestre Fotoservizi 📸. Hai usato i "
        "nostri servizi in passato e vorremmo restare in contatto via WhatsApp "
        "con sconti riservati e novità.\n\n"
        "Rispondi SI per iscriverti, oppure ignora questo messaggio per uscire "
        "dalla lista.\n\n"
        "In qualunque momento puoi disiscriverti rispondendo STOP. »"
    )
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    add_para(doc, "")
    add_para(doc, "3.2 — Stati possibili di un contatto", bold=True, size=12)
    add_table(doc,
              headers=["Stato (optInStatus)", "Significato", "Riceve promo?"],
              rows=[
                  ["pending", "Messaggio opt-in non ancora inviato o appena inviato, attesa risposta", "NO"],
                  ["yes", "Ha risposto SI oppure si è registrato in app con consenso marketing attivo", "SÌ"],
                  ["no", "Ha risposto STOP, oppure 30 giorni senza risposta al soft opt-in", "NO (permanente)"],
              ])

    add_para(doc, "")
    add_para(doc, "3.3 — Come registriamo le risposte SI / STOP", bold=True, size=12)
    add_para(doc,
             "Il soft opt-in viene mandato via WhatsApp (vedi sezione 5.6.1 per "
             "il motivo tecnico per cui NON usiamo push o email per questa fase). "
             "Con WhatsApp manual batch non c'è modo di intercettare "
             "automaticamente le risposte — solo la Cloud API a pagamento lo "
             "permetterebbe. Quindi il flusso pratico è:")
    add_bullet(doc, "L'operatore legge le risposte ricevute su WhatsApp "
                    "Business (telefono o WhatsApp Web)")
    add_bullet(doc, "Apre l'app Silvestre operatore → Crea Promozione → Tab 1")
    add_bullet(doc, "Nella sezione «Inbox risposte soft opt-in» della Tab 1, "
                    "vede l'elenco dei 🟡 In attesa")
    add_bullet(doc, "Cerca per nome o numero il cliente che ha risposto su "
                    "WhatsApp")
    add_bullet(doc, "Tap sul bottone rapido «✅ SI ricevuto» (verde) → il sistema "
                    "aggiorna optInStatus=yes in Firestore")
    add_bullet(doc, "In alternativa, tap «❌ STOP ricevuto» (rosso) → "
                    "optInStatus=no permanente")

    add_para(doc, "")
    add_para(doc, "Tempo richiesto: ~2-3 secondi per cliente. Per un batch tipico "
                  "di 100 messaggi inviati arrivano 10-15 risposte SI nei primi "
                  "3-5 giorni. L'operatore dedica circa 30-45 secondi al giorno "
                  "durante la campagna soft opt-in. Sostenibile.", italic=True)

    add_para(doc, "")
    add_para(doc, "3.4 — Auto-cleanup dei pending senza risposta", bold=True, size=12)
    add_para(doc,
             "Una Cloud Function schedulata (cron job giornaliero) controlla "
             "tutti i contatti con optInStatus=pending. Se il messaggio soft "
             "opt-in è stato inviato più di 30 giorni fa (optInSentAt) e non "
             "c'è risposta, viene automaticamente segnato come optInStatus=no "
             "definitivo. L'operatore non deve farlo manualmente. Questo evita "
             "che la coda 'pending' cresca all'infinito e mantiene pulita la "
             "lista marketing.")

    # ============================================================
    # 4. ARCHITETTURA TECNICA
    # ============================================================
    add_heading(doc, "4. Architettura tecnica", level=1)
    add_para(doc, "4.1 — Storage dei contatti", bold=True, size=12)
    add_bullet(doc, "Collection Firestore marketing_contacts (contatti CSV "
                    "importati una tantum)")
    add_bullet(doc, "Collection Firestore users (clienti registrati nell'app, "
                    "con campo acceptedMarketing)")
    add_bullet(doc, "Firestore Security Rules: read solo se l'utente ha role "
                    "= admin/staff (cioè l'operatore). I clienti normali non "
                    "possono leggere nessuna lista, anche dall'APK scaricato")

    add_para(doc, "")
    add_para(doc, "4.2 — I 3 canali di invio promo (gratuiti)", bold=True, size=12)
    add_table(doc,
              headers=["Canale", "Costo", "Pace realistico", "Limite"],
              rows=[
                  ["Push notification in-app (FCM)",
                   "0 € illimitato",
                   "Istantaneo a tutti",
                   "Solo a chi ha l'app installata + acceptedMarketing=true"],
                  ["Email (Brevo free tier)",
                   "0 € fino a 300/giorno",
                   "Quasi istantaneo",
                   "Solo a chi ha email + acceptedMarketing=true"],
                  ["WhatsApp manual batch",
                   "0 € sempre",
                   "50-100 invii/giorno (anti-ban WhatsApp)",
                   "Operatore preme «Invia» 1 volta per ogni contatto"],
              ])

    add_callout(doc, "Perché 3 canali",
                "FCM raggiunge solo chi ha l'app (oggi pochi, in crescita nel "
                "tempo). Email raggiunge chi ha lasciato l'indirizzo. WhatsApp "
                "raggiunge tutti gli altri ma è lento. Combinandoli "
                "intelligentemente copri quasi il 100% della base opted-in.")

    add_para(doc, "")
    add_para(doc, "4.3 — Perché NON usiamo WhatsApp Cloud API a pagamento", bold=True, size=12)
    add_para(doc,
             "La Cloud API ufficiale di Meta permetterebbe l'invio massivo "
             "automatico, ma costa circa 0,07 € per messaggio marketing in Italia. "
             "Per un singolo invio a 6000+ contatti si pagherebbero ~470 €, e "
             "richiederebbe l'approvazione preventiva di ogni template da parte "
             "di Meta (tempistica 1-3 giorni). Scelta esplicita del committente: "
             "rimanere su canali gratuiti, accettando il pace più lento.")

    # ============================================================
    # 5. IL PANNELLO «CREA PROMOZIONE»
    # ============================================================
    add_heading(doc, "5. Il pannello «Crea Promozione» nell'app operatore", level=1)
    add_para(doc, "5.1 — Struttura generale", bold=True, size=12)
    add_para(doc,
             "Nuova card nella dashboard operatore «Crea Promozione». "
             "Aprendola si arriva a una schermata con 4 tab in alto, in questo "
             "ordine preciso (la sequenza è progettata per ricordare le regole "
             "di compliance PRIMA che l'operatore inizi a creare promozioni):")
    add_bullet(doc, "Tab 1 — «Logica & GDPR» (compliance first, sempre come prima vista)")
    add_bullet(doc, "Tab 2 — «Crea Promozione» (form + invio)")
    add_bullet(doc, "Tab 3 — «Contatti Inclusi» (selezione destinatari)")
    add_bullet(doc, "Tab 4 — «Contatti Esclusi» (re-include)")

    add_para(doc, "")
    add_para(doc, "5.2 — Tab 1: Logica & Conformità GDPR", bold=True, size=12)
    add_para(doc,
             "PRIMA tab del pannello. Read-only. Mostra in modo trasparente la "
             "logica di gestione del consenso, lo stato corrente della base "
             "marketing e fornisce il punto di accesso alla campagna soft "
             "opt-in. Due scopi: (a) ricordare all'operatore le regole "
             "anti-spam del sistema PRIMA di creare promozioni, (b) dimostrare "
             "al Garante Privacy in caso di verifica che il software rispetta "
             "le scelte dei clienti.")

    add_para(doc, "")
    add_para(doc, "5.2.1 — Gli stati di un contatto (optInStatus)", bold=True, size=11)
    add_table(doc,
              headers=["Stato", "Significato",
                       "Riceve promo standard?", "Riceve soft opt-in?"],
              rows=[
                  ["🟢 yes",
                   "Ha dato consenso marketing (via app o risposta SI al WhatsApp)",
                   "Sì", "No (è già dentro)"],
                  ["⚪ pending mai contattato",
                   "Nel CSV ma soft opt-in mai inviato — «Nuovo»",
                   "No", "Sì (candidato)"],
                  ["🟡 pending già contattato",
                   "Soft opt-in inviato, attesa risposta SI/STOP — «In attesa»",
                   "No", "No (già contattato)"],
                  ["🔴 no",
                   "Ha risposto STOP o 30 giorni senza risposta",
                   "No mai", "No mai"],
              ])

    add_para(doc, "")
    add_para(doc, "5.2.2 — Le 4 categorie di cliente (regole di dedup)",
             bold=True, size=11)
    add_table(doc,
              headers=["Categoria", "Riceve soft opt-in?", "Motivo"],
              rows=[
                  ["A — Solo CSV, mai entrato in app",
                   "SÌ",
                   "Mai chiesto consenso, soft opt-in è richiesta legittima"],
                  ["B — Registrato in app con acceptedMarketing=true",
                   "NO",
                   "Già opted-in dall'app, è già in lista marketing attiva"],
                  ["C — Registrato in app con acceptedMarketing=false",
                   "NO ASSOLUTO",
                   "Ha esplicitamente rifiutato — GDPR vieta nuova richiesta"],
                  ["D — CSV + Registrato in app con marketing rifiutato",
                   "NO",
                   "La scelta esplicita più recente (in app) vince sul CSV"],
              ])

    add_para(doc, "")
    add_para(doc, "5.2.3 — La REGOLA D'ORO", bold=True, size=11)
    p = doc.add_paragraph()
    r = p.add_run(
        "La scelta esplicita del cliente vince SEMPRE sullo stato CSV."
    )
    r.bold = True
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xF4, 0x75, 0x21)

    add_para(doc, "Esempi pratici:")
    add_bullet(doc, "Cliente in CSV con pending → si registra in app spuntando "
                    "marketing → diventa yes. Da ora riceve le promo.")
    add_bullet(doc, "Cliente in CSV con pending → si registra in app SENZA "
                    "spuntare marketing → diventa no definitivo. Mai più contattato.")
    add_bullet(doc, "Cliente in CSV con pending → riceve soft opt-in → "
                    "risponde «SI» → diventa yes. Riceve le future promo.")
    add_bullet(doc, "Cliente in CSV con pending → riceve soft opt-in → "
                    "risponde «STOP» → diventa no definitivo.")
    add_bullet(doc, "Cliente in CSV con pending → riceve soft opt-in → non "
                    "risponde per 30 giorni → cron job lo segna no automatico.")

    add_para(doc, "")
    add_para(doc, "5.2.4 — Cosa NON deve mai succedere (impedito by design)",
             bold=True, size=11)
    add_bullet(doc, "Mandare promo a 🟡 In attesa o ⚪ Nuovo "
                    "(Tab 3 li filtra automaticamente, non li mostra proprio)")
    add_bullet(doc, "Mandare soft opt-in a chi ha rifiutato in app "
                    "(lo script di dedup li salta automaticamente)")
    add_bullet(doc, "Riprovare soft opt-in dopo che il cliente ha mandato STOP "
                    "(passa subito a no permanente)")
    add_bullet(doc, "Mandare soft opt-in due volte allo stesso «In attesa» "
                    "(i 🟡 non sono ricandidati al soft opt-in)")

    add_para(doc, "")
    add_para(doc, "5.2.5 — Auto-cleanup giornaliero", bold=True, size=11)
    add_para(doc,
             "Una Cloud Function schedulata gira ogni giorno alle 03:00 e:")
    add_bullet(doc, "Trova tutti i contatti con optInStatus=pending "
                    "e optInSentAt < oggi - 30 giorni")
    add_bullet(doc, "Li marca optInStatus=no definitivo")
    add_bullet(doc, "Logga l'azione in collection audit_log per tracciabilità "
                    "Garante Privacy")

    add_para(doc, "")
    add_para(doc, "5.2.6 — Statistiche live mostrate in cima alla Tab 1",
             bold=True, size=11)
    add_para(doc,
             "Pannello che si aggiorna in tempo reale dal Firestore. Esempio "
             "di rendering:")
    p = doc.add_paragraph()
    r = p.add_run(
        "🟢 Acconsentiti (yes):              1.847\n"
        "⚪ Nuovi mai contattati:            4.501\n"
        "🟡 In attesa risposta:                111\n"
        "🔴 Rifiutati/STOP/scaduti:            310\n"
        "                                  ──────\n"
        "📦 Totale contatti gestiti:        6.769\n"
        "📱 Di cui registrati in app:          234"
    )
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    add_para(doc, "")
    add_para(doc, "5.2.7 — Bottone «Lancia campagna soft opt-in»",
             bold=True, size=11)
    add_para(doc,
             "In fondo alla Tab 1, un riquadro CTA dedicato al soft opt-in:")
    p = doc.add_paragraph()
    r = p.add_run(
        "« HAI N CLIENTI NUOVI DA CONTATTARE\n\n"
        "Lancia la campagna di richiesta consenso (soft opt-in).\n"
        "Pace consigliato: 50-100 messaggi al giorno per non triggerare "
        "l'anti-spam di WhatsApp.\n\n"
        "[ 🚀 LANCIA / RIPRENDI CAMPAGNA SOFT OPT-IN ] »"
    )
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    add_para(doc, "")
    add_para(doc,
             "Cliccando, apre un workflow dedicato (vedi sezione 5.6) che "
             "lavora ESCLUSIVAMENTE su contatti ⚪ Nuovi (i 🟡 In attesa sono "
             "già stati contattati e non si ricontattano fino a scadenza dei "
             "30 giorni).")

    add_para(doc, "")
    add_para(doc, "5.2.8 — Storico campagne", bold=True, size=11)
    add_para(doc,
             "Tabella delle ultime 10 campagne (sia promo standard che soft "
             "opt-in): data, canale, numero destinatari, esito. Permette "
             "all'operatore di verificare quando ha mandato cosa e a chi.")

    add_para(doc, "")
    add_para(doc, "5.3 — Tab 2: Crea Promozione", bold=True, size=12)
    add_bullet(doc, "Campo «Titolo promozione» (testo breve)")
    add_bullet(doc, "Campo «Dettagli promozione» (testo multilinea)")
    add_bullet(doc, "Upload «Foto esempio» (max 3 immagini, salvate su Cloudinary)")
    add_bullet(doc, "Range validità: data picker «Dal» — data picker «Al»")
    add_bullet(doc, "Campo «Costo/Sconto» (testo libero, es. «-30%» o «10 €»)")
    add_bullet(doc, "Anteprima live del messaggio composto coi placeholder sostituiti")
    add_bullet(doc, "3 bottoni di invio nella sezione inferiore («Invia Push», "
                    "«Invia Email», «Invia WhatsApp»), tutti disabilitati finché "
                    "i campi obbligatori non sono compilati")

    add_para(doc, "")
    add_para(doc, "5.3.1 — Alert di conferma pre-invio", bold=True, size=11)
    add_para(doc,
             "Quando l'operatore clicca uno dei bottoni di invio, prima parte "
             "un dialog di conferma:")
    p = doc.add_paragraph()
    r = p.add_run(
        "« Stai per inviare la promozione a N contatti via {canale}.\n\n"
        "Hai impostato correttamente i contatti inclusi (Tab 3) ed esclusi "
        "(Tab 4)? »"
    )
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    add_para(doc, "")
    add_para(doc,
             "Due bottoni: SÌ, INVIA (verde) → procede con l'invio. "
             "NO, INDIETRO (grigio) → torna al pannello senza inviare.")

    add_para(doc, "")
    add_para(doc, "5.4 — Tab 3: Destinatari Promozione", bold=True, size=12)
    add_para(doc,
             "Lista dei clienti che possono ricevere la promo standard che "
             "stai creando in Tab 2. PRINCIPIO ANTI-ERRORE: questa tab "
             "contiene SOLO i clienti con consenso marketing attivo "
             "(🟢 Acconsentiti). I 🟡 In attesa e ⚪ Nuovi NON compaiono "
             "affatto in questa lista — sono gestiti separatamente dalla "
             "campagna soft opt-in (vedi 5.6). Conseguenza: l'operatore "
             "NON può sbagliare e includere per errore chi non ha "
             "acconsentito. È impossibile by design.")

    add_para(doc, "")
    add_para(doc, "Mockup della Tab 3:", italic=True)
    p = doc.add_paragraph()
    r = p.add_run(
        "🛡 Questa lista contiene SOLO i clienti che hanno dato consenso\n"
        "   marketing (N). I M in attesa o nuovi non sono qui — usali con\n"
        "   la campagna soft opt-in (Tab 1).\n\n"
        "🔍 [Cerca nome o telefono...]        [TUTTI] [📱 App] [📞 CSV]\n\n"
        "[✓ INCLUDI TUTTI N]  [✗ ESCLUDI TUTTI]   N inclusi / N acconsentiti\n\n"
        "[✓] 📱 Mario Rossi    +39 335 12345 67    [✓ INCLUDI] [✗ ESCLUDI]\n"
        "[✓] 📞 Anna Bianchi   +39 348 99 81 245   [✓ INCLUDI] [✗ ESCLUDI]\n"
        "[✓] 📱 Luca Verdi     +39 333 77 22 109   [✓ INCLUDI] [✗ ESCLUDI]\n"
        "..."
    )
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    add_para(doc, "")
    add_para(doc, "Componenti:")
    add_bullet(doc, "Banner in cima che spiega la lista filtrata e rimanda alla "
                    "campagna soft opt-in per gli altri")
    add_bullet(doc, "Barra di ricerca (cerca per nome / numero di telefono)")
    add_bullet(doc, "Filtri rapidi a chip: Tutti | 📱 Solo app | 📞 Solo CSV "
                    "(NON più 🟡 / ⚪ perché non esistono qui)")
    add_bullet(doc, "Bottone «INCLUDI TUTTI N» (verde) — sicuro by design "
                    "perché agisce solo sugli N acconsentiti")
    add_bullet(doc, "Bottone «ESCLUDI TUTTI» (rosso) — per ripartire da zero "
                    "selezione")
    add_bullet(doc, "Default all'apertura: tutti gli acconsentiti spuntati "
                    "(massimo opt-in del segmento)")
    add_bullet(doc, "Lista contatti scrollabile. Ogni riga contiene:")
    add_bullet(doc, "  ‣ Checkbox a sinistra (spuntata = incluso, vuota = escluso)",
               level=1)
    add_bullet(doc, "  ‣ Badge sorgente (📱 app / 📞 solo CSV)", level=1)
    add_bullet(doc, "  ‣ Nome e telefono del cliente", level=1)
    add_bullet(doc, "  ‣ Bottone «INCLUDI» (verde) — spunta checkbox", level=1)
    add_bullet(doc, "  ‣ Bottone «ESCLUDI» (rosso) — sposta in Tab 4", level=1)
    add_bullet(doc, "Contatore live in cima: «N inclusi / M acconsentiti totali»")

    add_para(doc, "")
    add_para(doc, "5.4.1 — Sanity check finale prima dell'invio",
             bold=True, size=11)
    add_para(doc,
             "Con Tab 3 ristretta agli acconsentiti by design, il vecchio "
             "«smart guardrail» diventa superfluo. Rimane un sanity check "
             "leggero: prima dell'invio finale, il sistema rilegge gli "
             "optInStatus dei selezionati da Firestore. Se qualcuno è passato "
             "a optInStatus=no tra la selezione e l'invio (caso edge: cliente "
             "ha mandato STOP nel frattempo), viene rimosso automaticamente "
             "dal batch e l'operatore vede un messaggio: «3 contatti rimossi "
             "perché hanno revocato il consenso dopo la tua selezione.»")

    add_para(doc, "")
    add_para(doc, "5.5 — Tab 4: Contatti Esclusi", bold=True, size=12)
    add_bullet(doc, "Stessa barra di ricerca della Tab 3")
    add_bullet(doc, "Lista contatti che l'operatore ha escluso manualmente da "
                    "questa specifica campagna (es. cliente che si lamenta di "
                    "ricevere troppe promo)")
    add_bullet(doc, "Bottone per riga: «RE-INCLUDI» (verde) → rimette in Tab 3")
    add_bullet(doc, "I contatti con optInStatus=no (STOP definitivo o no-reply "
                    "30 gg) NON compaiono qui — sono nascosti del tutto, mai "
                    "più raggiungibili")
    add_bullet(doc, "I 🟡 In attesa e ⚪ Nuovi NON compaiono qui (non sono mai "
                    "stati in Tab 3, quindi non possono essere stati esclusi)")

    add_para(doc, "")
    add_para(doc, "5.6 — Workflow «Campagna soft opt-in» (separato)",
             bold=True, size=12)
    add_para(doc,
             "Workflow dedicato, completamente separato dal pannello promo "
             "standard. Accessibile solo dal bottone «🚀 Lancia campagna soft "
             "opt-in» nella Tab 1. Esiste apposta separato per evitare ogni "
             "rischio di mischiare promo e richiesta consenso.")

    add_para(doc, "")
    add_para(doc, "5.6.1 — Canale di invio: SOLO WhatsApp manual batch",
             bold=True, size=11)
    add_para(doc,
             "Il messaggio soft opt-in viene inviato SOLO via WhatsApp manual "
             "batch. Motivo tecnico: i clienti in soft opt-in (⚪ Nuovi) sono "
             "per definizione NON registrati nell'app — se lo fossero, "
             "avrebbero già fatto la scelta marketing sì/no in registrazione. "
             "Pertanto:")
    add_table(doc,
              headers=["Canale", "Usato per soft opt-in?", "Motivo"],
              rows=[
                  ["Push notification (FCM)",
                   "NO",
                   "FCM raggiunge solo utenti app già registrati. I ⚪ Nuovi "
                   "non hanno mai installato l'app, irraggiungibili via FCM"],
                  ["Email (Brevo)",
                   "NO (per ora)",
                   "Possibile fallback futuro per chi ha email nel CSV ma "
                   "non risponde a WhatsApp. Non implementato nella prima fase"],
                  ["WhatsApp manual batch",
                   "SÌ — unico canale",
                   "Tutti i ⚪ Nuovi hanno numero di telefono. È il canale "
                   "diretto e gratuito per raggiungerli e ricevere la "
                   "risposta SI/STOP conversazionale"],
              ])

    add_para(doc, "")
    add_para(doc, "5.6.2 — Template fisso del messaggio soft opt-in",
             bold=True, size=11)
    add_para(doc,
             "Nel workflow soft opt-in il messaggio NON è modificabile "
             "dall'operatore (a differenza della promo standard, dove "
             "titolo/dettagli/costo sono liberi). Il testo è fisso e "
             "approvato:")
    p = doc.add_paragraph()
    r = p.add_run(
        "« Ciao {{nome}}, ti scriviamo da Silvestre Fotoservizi 📸. Hai "
        "usato i nostri servizi in passato e vorremmo restare in contatto "
        "via WhatsApp con sconti riservati e novità.\n\n"
        "Rispondi SI per iscriverti, oppure ignora questo messaggio per "
        "uscire dalla lista.\n\n"
        "In qualunque momento puoi disiscriverti rispondendo STOP. »"
    )
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    add_para(doc, "Bloccare il testo elimina il rischio che l'operatore "
                  "trasformi accidentalmente la richiesta consenso in una "
                  "promo mascherata (che sarebbe spam).", italic=True)

    add_para(doc, "")
    add_para(doc, "5.6.3 — Lista bersaglio: SOLO ⚪ Nuovi NON registrati in app",
             bold=True, size=11)
    add_bullet(doc, "Il workflow mostra unicamente contatti con "
                    "optInStatus=pending AND optInSentAt=null AND NESSUN "
                    "user registrato in app con telefono corrispondente")
    add_bullet(doc, "Chi ha scaricato l'app non è MAI candidato al soft "
                    "opt-in — ha già scelto sì/no marketing in registrazione")
    add_bullet(doc, "I 🟡 In attesa non compaiono — sono già stati contattati "
                    "una volta. Se non rispondono entro 30 giorni vengono "
                    "auto-segnati no (mai più richiamati)")
    add_bullet(doc, "Possibilità di filtrare per numero d'ordine recente "
                    "(es. «mostra solo chi ha ordinato negli ultimi 24 mesi»). "
                    "Più alto il rapporto di recency, più alta la probabilità "
                    "che il cliente ti riconosca e risponda")
    add_bullet(doc, "Rate-limit di sicurezza: max 100 invii al giorno "
                    "(consigliato 50-80 per non triggerare l'anti-spam di "
                    "WhatsApp)")

    add_para(doc, "")
    add_para(doc, "5.6.4 — Avanzamento e tracking", bold=True, size=11)
    add_bullet(doc, "Ogni invio aggiorna optInSentAt=timestamp(now) e sposta "
                    "il contatto da ⚪ Nuovo a 🟡 In attesa")
    add_bullet(doc, "Progress bar in cima: «12 / 50 inviati oggi»")
    add_bullet(doc, "Stato persistito in Firestore: se chiudi l'app, il "
                    "giorno dopo riprendi da dove eri rimasto")
    add_bullet(doc, "Per registrare la risposta del cliente vedi sezione 3.3")

    # ============================================================
    # 6. FLUSSO OPERATIVO TIPICO
    # ============================================================
    add_heading(doc, "6. Flusso operativo tipico (esempio reale)", level=1)
    add_para(doc, "Esempio: l'operatore vuole mandare la promo «Stampe in Tela "
                  "-30% dal 1 al 15 giugno».", italic=True)
    add_bullet(doc, "Apre l'app operatore → tap su card «Crea Promozione»")
    add_bullet(doc, "Atterra su Tab 1 «Logica & GDPR»: dà un'occhiata alle "
                    "statistiche live (es. 1.847 opted-in, 4.612 pending) e si "
                    "ricorda le regole prima di procedere")
    add_bullet(doc, "Tab 2: compila titolo «Stampe Tela -30%», dettagli, foto "
                    "di una stampa su tela, date 01/06 — 15/06, costo «-30%»")
    add_bullet(doc, "Vede l'anteprima del messaggio. Vuole controllare le liste.")
    add_bullet(doc, "Tab 3: vede i contatti opted-in (es. 1.847). Toglie la "
                    "checkbox a 3 clienti che lo hanno chiamato lamentandosi "
                    "delle troppe promo → vanno automaticamente in Tab 4.")
    add_bullet(doc, "Torna Tab 2. Vede contatore aggiornato: «1.844 contatti».")
    add_bullet(doc, "Clicca «Invia via Push» → arriva l'alert «Stai per inviare "
                    "a 1.844 contatti via Push» → conferma «SÌ, INVIA»")
    add_bullet(doc, "Le push partono istantaneamente a tutti i 1.844 utenti app "
                    "che hanno marketing attivo")
    add_bullet(doc, "Decide di inviare anche via WhatsApp: clicca «Invia "
                    "WhatsApp» → conferma → l'app apre WhatsApp con il primo "
                    "contatto precompilato. Operatore preme Invia su WhatsApp, "
                    "torna in app, parte automaticamente al secondo. Pace ~50 "
                    "invii in mezz'ora.")
    add_bullet(doc, "Il giorno dopo riprende il batch da dove si era fermato "
                    "(progresso salvato in Firestore). In 30-40 giorni copre i "
                    "1.844 contatti opted-in via WhatsApp.")

    # ============================================================
    # 7. SICUREZZA DEI DATI
    # ============================================================
    add_heading(doc, "7. Sicurezza e privacy operativa", level=1)
    add_bullet(doc, "I 6000+ contatti del CSV vengono importati una sola volta "
                    "in Firestore (collection marketing_contacts), poi il file "
                    "fisico .csv non viene mai più toccato")
    add_bullet(doc, "Firestore Security Rules: «allow read: if request.auth != "
                    "null && request.auth.token.role == 'admin'». Nessun cliente "
                    "normale può vedere la lista.")
    add_bullet(doc, "I dati sensibili (telefoni, email) non vengono mai "
                    "bundlizzati nel build dell'app. Restano solo lato server.")
    add_bullet(doc, "Anche se qualcuno reverse-engineering il codice client, "
                    "non trova mai la lista contatti.")
    add_bullet(doc, "Audit trail: ogni invio promozionale registra in Firestore "
                    "(collection promotions_sent) chi è stato contattato, "
                    "quando, su che canale. Serve per dimostrare al Garante "
                    "che si rispetta il consenso.")

    # ============================================================
    # 8. ROADMAP DI IMPLEMENTAZIONE
    # ============================================================
    add_heading(doc, "8. Roadmap di implementazione", level=1)
    add_table(doc,
              headers=["#", "Task", "Stima", "Stato"],
              rows=[
                  ["1", "Re-introdurre Satispay nel checkout", "30 min", "✓ FATTO"],
                  ["2", "Aggiungere Bonifico Istantaneo come metodo pagamento", "1 h", "In attesa IBAN"],
                  ["3", "Box marketing evidenziata in registrazione + Impostazioni", "30 min", "Prossimo"],
                  ["4", "Import CSV 6000+ contatti in Firestore + rules", "1 h", "Pianificato"],
                  ["5", "Pannello «Crea Promozione» 4 tab (3 operative + Logica/GDPR) + 3 canali invio", "6-7 h", "Pianificato"],
                  ["6", "Campagna soft opt-in WhatsApp (manual batch)", "Operatore: 60-90 gg", "Dopo task 5"],
              ])

    # ============================================================
    # 9. DECISIONI PENDENTI
    # ============================================================
    add_heading(doc, "9. Decisioni pendenti dal committente", level=1)
    add_bullet(doc, "IBAN, intestatario e banca per il bonifico istantaneo "
                    "(serve per task 2)")
    add_bullet(doc, "QR in negozio per acquisire nuovi utenti app: "
                    "soprasseduto per il momento, eventualmente da aggiungere "
                    "in seconda battuta")
    add_bullet(doc, "Integrazione Brevo (email broadcast): credenziali API "
                    "Brevo da generare quando si attiva il canale email")
    add_bullet(doc, "Verifica volume effettivo email valide nel CSV (molti "
                    "contatti hanno solo telefono, l'email è vuota)")

    # ============================================================
    # 10. APPENDICE — DATI CHIAVE
    # ============================================================
    add_heading(doc, "10. Appendice — Dati chiave e riferimenti", level=1)
    add_table(doc,
              headers=["Dato", "Valore"],
              rows=[
                  ["Contatti CSV totali", "6000+ contatti"],
                  ["Numero WhatsApp Business negozio", "+39 335 169 7903"],
                  ["Email negozio", "fotosilvestre1970@gmail.com"],
                  ["Indirizzo negozio", "Via V. Emanuele III, 205 — Frattamaggiore (NA)"],
                  ["Firebase project", "silvestre-fotoservizi-it"],
                  ["File CSV originale", "Desktop/SilvestreApp/Contatti_Clienti.csv"],
                  ["Provider pagamenti scelto", "Satispay (online) + Bonifico Istantaneo + Caparra in negozio"],
                  ["Provider email (gratis)", "Brevo (ex Sendinblue) — 300 email/giorno free tier"],
                  ["Provider push (gratis)", "Firebase Cloud Messaging — illimitato"],
              ])

    # ============================================================
    # FOOTER
    # ============================================================
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Documento generato il 20 maggio 2026 — Silvestre Fotoservizi · "
        "Frattamaggiore (NA)"
    )
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    out_path = (
        Path(__file__).resolve().parent.parent / "docs" / "Promozioni.docx"
    )
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"[OK] Documento generato: {out_path}")
    print(f"     Pagine stimate: ~6-8 A4")


if __name__ == "__main__":
    main()
